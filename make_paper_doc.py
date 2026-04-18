"""
Academic-paper-style Word document — 全中文版
包含：系統架構、SECOM流程圖、製程流程圖、名詞說明
"""
import io
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
from matplotlib.patches import FancyBboxPatch
import numpy as np

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── matplotlib 中文字型 ────────────────────────────────────────────────────────
plt.rcParams['font.family'] = ['Microsoft JhengHei', 'Microsoft YaHei',
                                'DFKai-SB', 'MingLiU', 'DejaVu Sans']
plt.rcParams['axes.unicode_minus'] = False

# ── 顏色 ──────────────────────────────────────────────────────────────────────
NAVY   = '#1A3A5C'
BLUE   = '#2E86AB'
GREEN  = '#1B7340'
ORANGE = '#E67E22'
RED    = '#C0392B'
PURPLE = '#6C3483'
TEAL   = '#0D6E6E'
GRAY   = '#555555'
LGRAY  = '#AAAAAA'
LLGRAY = '#F0F0F0'

def dk(hexc, pct=0.6):
    r,g,b = int(hexc[1:3],16), int(hexc[3:5],16), int(hexc[5:7],16)
    return f'#{int(r*pct):02X}{int(g*pct):02X}{int(b*pct):02X}'

def lt(hexc, mix=0.85):
    r,g,b = int(hexc[1:3],16), int(hexc[3:5],16), int(hexc[5:7],16)
    r2 = int(r + (255-r)*mix); g2 = int(g + (255-g)*mix); b2 = int(b + (255-b)*mix)
    return f'#{r2:02X}{g2:02X}{b2:02X}'

# ── 通用繪圖 helpers ───────────────────────────────────────────────────────────
def fig_buf(fig):
    buf = io.BytesIO()
    fig.savefig(buf, format='png', dpi=180, bbox_inches='tight',
                facecolor=fig.get_facecolor())
    buf.seek(0); plt.close(fig); return buf

def box(ax, x, y, w, h, text, sub='', fc=BLUE, ec=None, fs=9, tcolor='white', bold=True):
    ec = ec or dk(fc)
    r = FancyBboxPatch((x,y),w,h, boxstyle='round,pad=0.06',
                       lw=1.5, edgecolor=ec, facecolor=fc, zorder=2)
    ax.add_patch(r)
    if sub:
        ax.text(x+w/2, y+h*0.68, text, ha='center', va='center',
                fontsize=fs, fontweight='bold' if bold else 'normal',
                color=tcolor, zorder=3, multialignment='center')
        ax.text(x+w/2, y+h*0.28, sub, ha='center', va='center',
                fontsize=fs-1.5, color=tcolor, zorder=3,
                multialignment='center', alpha=0.88)
    else:
        ax.text(x+w/2, y+h/2, text, ha='center', va='center',
                fontsize=fs, fontweight='bold' if bold else 'normal',
                color=tcolor, zorder=3, multialignment='center')

def diamond(ax, cx, cy, hw, hh, text, fc='#F0B429', ec='#C78B00', fs=8.5):
    pts = np.array([[cx,cy+hh],[cx+hw,cy],[cx,cy-hh],[cx-hw,cy]])
    ax.add_patch(plt.Polygon(pts, fc=fc, ec=ec, lw=1.8, zorder=2))
    ax.text(cx, cy, text, ha='center', va='center',
            fontsize=fs, fontweight='bold', color='#1A1A1A', zorder=3,
            multialignment='center')

def arr(ax, x1,y1,x2,y2, label='', c=GRAY, lw=1.8, style='->',
        rad=0.0, lpos='mid', loff=(0.08,0.12)):
    cs = f'arc3,rad={rad}' if rad else 'arc3,rad=0'
    ax.annotate('', xy=(x2,y2), xytext=(x1,y1),
                arrowprops=dict(arrowstyle=style, color=c, lw=lw,
                                connectionstyle=cs), zorder=1)
    if label:
        mx = (x1+x2)/2 + loff[0]; my = (y1+y2)/2 + loff[1]
        ax.text(mx, my, label, fontsize=7.5, color=c, ha='left', va='bottom')

# ══════════════════════════════════════════════════════════════════════════════
# 圖 0：論文用系統整體架構圖（載入 make_paper_arch.py 輸出之 PNG）
# ══════════════════════════════════════════════════════════════════════════════
def fig_arch_paper():
    path = r"c:\Users\user\Desktop\在職碩\OneDrive - 長庚大學\長庚碩班\論文\系統架構圖_論文用.png"
    with open(path, 'rb') as f:
        buf = io.BytesIO(f.read())
    return buf

# ══════════════════════════════════════════════════════════════════════════════
# 圖 1：系統架構分層圖
# ══════════════════════════════════════════════════════════════════════════════
def fig_arch():
    fig, ax = plt.subplots(figsize=(13,7.5))
    fig.patch.set_facecolor('#F8F9FA'); ax.set_facecolor('#F8F9FA')
    ax.set_xlim(0,10); ax.set_ylim(0,7.5); ax.axis('off')

    layers = [
        (6.6, BLUE,   '第一層：使用者介面層',
         'viewer.html（Three.js WebGL）  ·  PointerLock 第一人稱視角  ·  HMI 虛擬螢幕  ·  Chat 對話框  ·  CDU/Overlay 監控'),
        (5.5, GREEN,  '第二層：REST API 層（port 8765）',
         'GET /api/hmi  /api/wafer_map  /api/status  /api/process_window  /api/secom_summary\n'
         'POST /api/start  /api/chat  /api/exposure  /api/fault'),
        (4.4, ORANGE, '第三層：業務邏輯層',
         'SimulationTrainingSystem  ·  ScenarioEngine  ·  NLU 自然語言控制器  ·  ClosedLoopController  ·  AIScenarioMentor'),
        (3.3, RED,    '第四層：物理模型層',
         'LensHeatingEngine（雙指數熱模型）  ·  Bossung Curve（CD vs 焦距/劑量）  ·  Zernike 像差（Z4/Z7/Z9）  ·  CDU Map 13×13'),
        (2.2, PURPLE, '第五層：資料層',
         'SecomNoiseModel  ·  UCI SECOM（1,567 samples，590 features）  ·  LithographyDigitalTwin  ·  ProcessParameterDB'),
        (1.1, TEAL,   '第六層：3D 資產層',
         'asml_duv.glb（ASML DUV 設備模型）  ·  GLTFLoader.js  ·  Mesh 互動  ·  Material 高亮  ·  外殼/內部雙模式'),
    ]
    for y, col, label, content in layers:
        r = FancyBboxPatch((0.12,y-0.48),9.76,0.88,
                           boxstyle='round,pad=0.05', lw=1.5,
                           edgecolor=col, facecolor=lt(col,0.88))
        ax.add_patch(r)
        badge = FancyBboxPatch((0.12,y-0.48),2.7,0.88,
                               boxstyle='round,pad=0.05', lw=0,
                               facecolor=col, zorder=2)
        ax.add_patch(badge)
        ax.text(1.47, y, label, ha='center', va='center',
                fontsize=8.5, fontweight='bold', color='white', zorder=3,
                multialignment='center')
        ax.text(5.3, y, content, ha='center', va='center',
                fontsize=8, color='#2C2C2C', zorder=3, multialignment='center')
        if y > 1.1:
            ax.annotate('', xy=(5.0,y-0.5), xytext=(5.0,y-0.62),
                        arrowprops=dict(arrowstyle='->', color='#888', lw=1.5))

    ax.set_title('圖 1　系統架構分層圖', fontsize=13, fontweight='bold',
                 pad=10, color=NAVY)
    fig.tight_layout(pad=0.4)
    return fig_buf(fig)

# ══════════════════════════════════════════════════════════════════════════════
# 圖 2：SECOM → 物理模型 → 輸出 流程圖（由上往下）
# ══════════════════════════════════════════════════════════════════════════════
def fig_secom_flow():
    fig, ax = plt.subplots(figsize=(10, 14))
    fig.patch.set_facecolor('#F8F9FA'); ax.set_facecolor('#F8F9FA')
    ax.set_xlim(0,10); ax.set_ylim(0,14); ax.axis('off')

    # ─ 由上到下 y 座標 ─
    # UCI SECOM Dataset
    box(ax, 2.5,12.3, 5.0,1.3,
        'UCI SECOM 資料集', '1,567 wafer samples  ·  590 sensor features\nPass: 1,463  /  Fail: 104',
        fc=PURPLE, fs=10)
    arr(ax, 5.0,12.3, 5.0,11.85, c=PURPLE, lw=2)

    # PCA 特徵分析
    box(ax, 2.5,10.7, 5.0,1.1,
        'PCA 特徵分析', 'Top-20 相關特徵萃取  ·  偏態分布建模（Skewness）',
        fc='#5B2C6F', fs=9.5)
    arr(ax, 5.0,10.7, 5.0,10.25, c=PURPLE, lw=2)

    # 健康分數計算
    box(ax, 2.5,9.1, 5.0,1.1,
        '設備健康分數計算', 'health_score ∈ [0, 1]  ·  分數越低代表設備狀態越差',
        fc='#76448A', fs=9.5)
    arr(ax, 5.0,9.1, 5.0,8.65, c=PURPLE, lw=2)

    # 故障注入判斷
    diamond(ax, 5.0,8.1, 2.2,0.58,
            '故障注入？\n(/api/fault)',
            fc='#F9E4F0', ec='#922B21')
    # Yes 右邊
    ax.annotate('', xy=(8.5,8.1), xytext=(7.2,8.1),
                arrowprops=dict(arrowstyle='->', color=RED, lw=1.8))
    ax.text(7.82,8.22,'是', fontsize=8.5, color=RED, fontweight='bold')
    box(ax, 8.5,7.7, 1.3,0.82,'故障類型注入',
        'dose_drift\nfocus_drift\nlens_hotspot\ncontamination\nstage_error',
        fc=RED, fs=7)
    ax.annotate('', xy=(9.15,7.7), xytext=(9.15,7.1),
                arrowprops=dict(arrowstyle='->', color=RED, lw=1.5,
                                connectionstyle='arc3,rad=0'))
    ax.annotate('', xy=(6.5,6.7), xytext=(9.15,7.1),
                arrowprops=dict(arrowstyle='->', color=RED, lw=1.5,
                                connectionstyle='arc3,rad=-0.3'))
    # No 往下
    arr(ax, 5.0,7.52, 5.0,7.1, '否', c=GREEN, lw=2, loff=(0.08,0.05))

    # noise_sigma 計算
    box(ax, 2.5,5.9, 5.0,1.1,
        'CD 雜訊標準差計算',
        'Normal: σ = 1.5 nm  ·  Fault: σ = 3.0 nm\n依 health_score 線性插值',
        fc=GREEN, fs=9.5)
    ax.text(7.6,6.45,'noise_σ 輸入', fontsize=8, color=GREEN, style='italic')
    arr(ax, 5.0,5.9, 5.0,5.45, c=GREEN, lw=2)

    # ── 製程參數輸入（左側） ──
    box(ax, 0.15,5.0, 1.85,0.9,'製程參數輸入',
        'dose, focus\nNA, sigma',
        fc=BLUE, fs=8.5)
    ax.annotate('', xy=(2.5,5.5), xytext=(2.0,5.45),
                arrowprops=dict(arrowstyle='->', color=BLUE, lw=1.8))
    ax.text(1.45,5.58,'參數', fontsize=8, color=BLUE)

    # LensHeatingEngine
    box(ax, 2.5,4.2, 5.0,1.15,
        'LensHeatingEngine  鏡片熱模型',
        'W(t) = A₁(1-e^(-t/τ₁)) + A₂(1-e^(-t/τ₂))\n5 鏡片元件：IL1, IL2, PL1, PL2, PL3  ·  ΔT → 焦距偏移',
        fc=RED, fs=9.5)
    arr(ax, 5.0,4.2, 5.0,3.75, c=RED, lw=2)

    # Zernike
    box(ax, 2.5,2.6, 5.0,1.1,
        'Zernike 像差計算',
        'Z4（散焦）  ·  Z7/Z8（彗差→Overlay偏移）  ·  Z9（球差）\nBossung Curve: CD = CD₀ + a·Δf² + b·Δdose',
        fc=ORANGE, fs=9.5)
    arr(ax, 5.0,2.6, 5.0,2.15, c=ORANGE, lw=2)

    # 輸出
    box(ax, 1.5,0.5, 7.0,1.55,
        '模擬輸出結果',
        'CD mean (nm)  ·  CD 3σ  ·  Overlay X/Y 3σ (nm)\n'
        'CDU Map 13×13  ·  Overlay Vector Field 7×7\n'
        '鏡片溫度  ·  製程良率預測  ·  Pass/Fail 判斷',
        fc=TEAL, fs=9.5)

    ax.set_title('圖 2　SECOM 資料驅動模擬流程圖', fontsize=13,
                 fontweight='bold', pad=10, color=NAVY)
    fig.tight_layout(pad=0.5)
    return fig_buf(fig)

# ══════════════════════════════════════════════════════════════════════════════
# 圖 3：微影製程流程圖
# ══════════════════════════════════════════════════════════════════════════════
def fig_process():
    fig, ax = plt.subplots(figsize=(13, 8))
    fig.patch.set_facecolor('#F8F9FA'); ax.set_facecolor('#F8F9FA')
    ax.set_xlim(0,13); ax.set_ylim(0,8); ax.axis('off')

    # 主流程（由左到右，y=5.5）
    steps_main = [
        (0.3,  '基板清洗\nWafer Clean',       '#1E3A5F'),
        (1.85, '底層薄膜沉積\nHard Mask / ARC', '#2C5F8A'),
        (3.4,  '光阻塗佈\nPR Coating (BARC)', '#1B5E20'),
        (4.95, '軟烘烤\nSoft Bake (PAB)',     '#1B5E20'),
        (6.5,  '曝光\nExposure (DUV 248nm)', RED),
        (8.05, '曝後烘烤\nPost Exposure Bake','#7D6608'),
        (9.6,  '顯影\nDevelop',              '#4A235A'),
        (11.15,'檢測\nInspection / OCD',     '#0D5C63'),
    ]
    BW = 1.35; BH = 1.0; BY = 5.5
    for x, label, col in steps_main:
        box(ax, x, BY, BW, BH, label, fc=col, fs=8.5)
        if x < 11.15:
            arr(ax, x+BW, BY+BH/2, x+BW+0.15, BY+BH/2, c='#444', lw=1.8)

    # PASS/FAIL 分支（y=5.5 → 下方）
    arr(ax, 11.15+BW, BY+BH/2, 12.3, BY+BH/2, c='#444', lw=1.8)
    diamond(ax, 12.65, BY+BH/2, 0.58,0.42, 'CD OK?',
            fc='#FEF5E7', ec='#E67E22')
    # Pass → 蝕刻
    arr(ax, 12.65, BY+BH/2-0.42, 12.65, 4.4, '通過', c=GREEN, lw=1.8, loff=(0.08,0.05))
    box(ax, 11.8, 3.4, 1.7, 0.9, '蝕刻 / 離子佈植\nEtch / Implant', fc='#1B4F72', fs=8)
    arr(ax, 12.65, 3.4, 12.65, 2.85, c='#1B4F72', lw=1.8)
    box(ax, 11.8, 1.95, 1.7, 0.8, '光阻去除\nStrip / Ash', fc='#154360', fs=8.5)
    arr(ax, 12.65, 1.95, 12.65, 1.4, c='#444', lw=1.8)
    box(ax, 11.8, 0.5, 1.7, 0.8, '下一道製程\nNext Layer', fc='#0B3D26', fs=8.5)

    # Fail → 重工
    ax.annotate('', xy=(9.6+BW/2, BY), xytext=(12.65, BY+BH/2-0.42),
                arrowprops=dict(arrowstyle='->', color=RED, lw=1.8,
                                connectionstyle='arc3,rad=-0.4'))
    ax.text(11.5, 4.9, '不合格\n→ 重工', fontsize=8, color=RED,
            fontweight='bold', ha='center')

    # 曝光站放大說明（y=1~4.5，x=0~7）
    rect_bg = FancyBboxPatch((0.1, 0.3), 11.4, 4.1,
                             boxstyle='round,pad=0.1', lw=1.5,
                             edgecolor=RED, facecolor=lt(RED,0.94))
    ax.add_patch(rect_bg)
    ax.text(5.8, 4.25, '曝光站（DUV ArF 248nm）詳細流程', ha='center',
            fontsize=10, fontweight='bold', color=RED)

    exp_steps = [
        (0.3,  2.2, '雷射能量\n設定\nDose Setting',    BLUE),
        (1.65, 2.2, '照明模式\n設定\nIllumination',    '#2C5F8A'),
        (3.0,  2.2, '光罩對準\nReticle\nAlignment',   GREEN),
        (4.35, 2.2, 'Wafer\n對準\nWafer Align',      '#1B5E20'),
        (5.7,  2.2, '焦距調整\nFocus / Level\nMapping', ORANGE),
        (7.05, 2.2, 'Step &\nScan 曝光\nExposure',    RED),
        (8.4,  2.2, '鏡片熱\n補償\nLens Heating',     '#7B241C'),
        (9.75, 2.2, 'APC\n製程控制\nAPC/FeedFwd',     PURPLE),
    ]
    EW=1.2; EH=1.5; EY=2.2
    for x, y, lbl, col in exp_steps:
        box(ax, x, y, EW, EH, lbl, fc=col, fs=7.5)
        if x < 9.75:
            arr(ax, x+EW, y+EH/2, x+EW+0.05, y+EH/2, c='#444', lw=1.5)

    # 參數監控欄
    monitors = [
        (0.3,  0.45, 'Dose Monitor\n(mJ/cm²)',   BLUE),
        (2.05, 0.45, 'Focus Monitor\n(nm)',        ORANGE),
        (3.8,  0.45, 'Overlay Monitor\n(nm)',      RED),
        (5.55, 0.45, 'CD-SEM\n量測 (nm)',           GREEN),
        (7.3,  0.45, 'Lens Temp\nMonitor (°C)',    '#7B241C'),
        (9.05, 0.45, 'SECOM\n製程資料',            PURPLE),
    ]
    for x, y, lbl, col in monitors:
        box(ax, x, y, 1.6, 0.7, lbl, fc=col, fs=7.8)

    ax.set_title('圖 4　微影製程流程圖（DUV 曝光站詳細）',
                 fontsize=13, fontweight='bold', pad=8, color=NAVY)
    fig.tight_layout(pad=0.4)
    return fig_buf(fig)

# ══════════════════════════════════════════════════════════════════════════════
# 圖 4：訓練系統流程圖
# ══════════════════════════════════════════════════════════════════════════════
def fig_training():
    fig, ax = plt.subplots(figsize=(13,6))
    fig.patch.set_facecolor('#F8F9FA'); ax.set_facecolor('#F8F9FA')
    ax.set_xlim(0,13); ax.set_ylim(0,6); ax.axis('off')

    def bx(x,y,w,h,t,items,fc,ec=None):
        ec = ec or dk(fc)
        r = FancyBboxPatch((x,y),w,h,boxstyle='round,pad=0.1',
                           lw=1.5,edgecolor=ec,facecolor=lt(fc,0.15))
        ax.add_patch(r)
        hdr2 = FancyBboxPatch((x,y+h-0.44),w,0.44,boxstyle='round,pad=0.04',
                              lw=0,facecolor=fc)
        ax.add_patch(hdr2)
        ax.text(x+w/2, y+h-0.22, t, ha='center', va='center',
                fontsize=9.5, fontweight='bold', color='white')
        for i, it in enumerate(items):
            ax.text(x+0.15, y+h-0.65-i*0.4, it, ha='left', va='top',
                    fontsize=8, color='#1A1A1A')

    # Start
    s = plt.Circle((0.5,3.0),0.32, fc=NAVY, ec=dk(NAVY), lw=2, zorder=2)
    ax.add_patch(s)
    ax.text(0.5,3.0,'開始', ha='center', va='center', fontsize=8,
            color='white', fontweight='bold', zorder=3)
    arr(ax,0.82,3.0,1.0,3.0,c=NAVY,lw=2)

    # 外框：整合式訓練情境
    outer = FancyBboxPatch((1.0,0.82),8.2,4.73,boxstyle='round,pad=0.1',
                           lw=2,edgecolor=dk(BLUE),facecolor=lt(BLUE,0.04))
    ax.add_patch(outer)
    hdr_outer = FancyBboxPatch((1.0,5.07),8.2,0.48,boxstyle='round,pad=0.04',
                               lw=0,facecolor=BLUE)
    ax.add_patch(hdr_outer)
    ax.text(5.1,5.31,'整合式訓練情境（理論問答  ×  3D 實機操作  同步進行）',
            ha='center',va='center',fontsize=10,fontweight='bold',color='white')

    # 左子框：AI 理論問答
    bx(1.2,0.98,3.2,3.85,'AI 理論問答',
       ['• Qwen LLM 蘇格拉底引導',
        '• 自適應教學模式切換',
        '• 挑戰/標準/鷹架/補救',
        '• 故障原因理論說明',
        '• 即時反饋與學習記錄'],
       BLUE)

    # 中間：雙向箭頭 + 標籤
    ax.annotate('',xy=(5.6,3.1),xytext=(4.4,3.1),
                arrowprops=dict(arrowstyle='<->',color='#555555',lw=2.0))
    ax.text(5.0,2.72,'AI 導師\n動態切換',ha='center',va='center',
            fontsize=8,color='#333',fontweight='bold',multialignment='center')

    # 右子框：3D 實機操作
    bx(5.8,0.98,3.2,3.85,'3D 實機操作',
       ['• 第一人稱 3D 環境',
        '• WASD 移動 + [E] 互動',
        '• 故障情境診斷（5 種）',
        '• 自主零件操作判斷',
        '• 求助學長（-5 分）'],
       GREEN)

    arr(ax,9.2,3.0,9.6,3.0,c=NAVY,lw=2)

    diamond(ax,10.15,3.0,0.58,0.45,'分數\n≥80?',fc='#FEF5E7',ec=ORANGE)
    ax.annotate('',xy=(6.4,0.82),xytext=(10.15,2.55),
                arrowprops=dict(arrowstyle='->',color=RED,lw=1.8,
                                connectionstyle='arc3,rad=0.4'))
    ax.text(8.7,0.58,'不合格，重試', fontsize=8, color=RED, fontweight='bold', ha='center')
    arr(ax,10.73,3.0,11.1,3.0,'合格',c=GREEN,lw=2,loff=(0.05,0.08))

    bx(11.1,1.3,1.7,3.4,'完成',
       ['• 綜合評分','• 能力認證','• 進階推薦'],
       TEAL)

    # Score bar
    score_rect = FancyBboxPatch((1.0,0.08),11.8,0.62,
                                boxstyle='round,pad=0.08',lw=1.2,
                                edgecolor=BLUE,facecolor=lt(BLUE,0.9))
    ax.add_patch(score_rect)
    ax.text(6.9,0.39,
            '評分：診斷準確性 30%  ｜  操作正確性 40%  ｜  安全合規 30%  ｜  時間效率加分',
            ha='center',va='center',fontsize=9.5,color=NAVY,fontweight='bold')

    ax.set_title('圖 7　訓練系統流程圖', fontsize=13,
                 fontweight='bold', pad=8, color=NAVY)
    fig.tight_layout(pad=0.4)
    return fig_buf(fig)

# ══════════════════════════════════════════════════════════════════════════════
# 圖 6：微影製程步驟圖（白底黑字學術風格，由上往下）
# ══════════════════════════════════════════════════════════════════════════════
def fig_process_steps():
    from matplotlib.patches import FancyBboxPatch, FancyArrowPatch
    fig, ax = plt.subplots(figsize=(10, 18))
    fig.patch.set_facecolor('white')
    ax.set_facecolor('white')
    ax.set_xlim(0, 10)
    ax.set_ylim(0, 18)
    ax.axis('off')

    BX, BW, BH = 2.0, 6.0, 1.05  # 主步驟 box 的 x, 寬, 高
    GAP = 0.45                     # 步驟間距
    EC = '#2C2C2C'
    TC = '#1A1A1A'

    def sbox(y, num, title, params='', dark=False):
        fc = '#2C2C2C' if dark else 'white'
        tc = 'white'   if dark else TC
        r = FancyBboxPatch((BX, y), BW, BH,
                           boxstyle='square,pad=0', lw=1.3,
                           edgecolor=EC, facecolor=fc, zorder=2)
        ax.add_patch(r)
        # 左側編號標籤
        nb = FancyBboxPatch((BX, y), 0.72, BH,
                            boxstyle='square,pad=0', lw=0,
                            facecolor='#2C2C2C', zorder=3)
        ax.add_patch(nb)
        ax.text(BX + 0.36, y + BH/2, num, ha='center', va='center',
                fontsize=10, fontweight='bold', color='white', zorder=4)
        # 步驟名稱
        ax.text(BX + 0.9, y + BH*0.68, title, ha='left', va='center',
                fontsize=10, fontweight='bold', color=TC, zorder=4)
        if params:
            ax.text(BX + 0.9, y + BH*0.28, params, ha='left', va='center',
                    fontsize=8.5, color='#444444', zorder=4, style='italic')

    def darr(y):
        ax.annotate('', xy=(5.0, y), xytext=(5.0, y + GAP),
                    arrowprops=dict(arrowstyle='->', color='#333333',
                                   lw=1.5, mutation_scale=14), zorder=1)

    steps = [
        ('①', '晶圓基板準備  Wafer Preparation',
         '矽晶圓切割、清洗、表面活化處理'),
        ('②', '底層薄膜沉積  Hard Mask / ARC',
         'CVD/PVD 沉積硬光罩或抗反射塗層（ARC）'),
        ('③', '光阻塗佈  Photoresist Coating',
         '旋塗 BARC + CAR 化學放大光阻，膜厚均勻性 ±0.5 nm'),
        ('④', '軟烘烤  Soft Bake（PAB）',
         '90–110 °C，60–90 s；驅除殘留溶劑，穩定膜層'),
        ('⑤', '晶圓對準  Wafer Alignment',
         '全域對準（Global）+ 精細對準（Fine），Overlay < 2 nm'),
        ('⑥', 'Step & Scan 曝光  DUV 248 nm Exposure',
         'Dose 設定、焦距調整（Focus/Level Map）、鏡片熱補償'),
        ('⑦', '曝後烘烤  Post-Exposure Bake（PEB）',
         '100–130 °C；酸催化反應、擴散改善站立波效應'),
        ('⑧', '顯影  Development',
         'TMAH 0.26N 顯影液，正型光阻溶解曝光區域'),
        ('⑨', '圖案檢測  CD Measurement / Inspection',
         'CD-SEM 量測線寬、OCD 量測膜厚、APC 回饋控制'),
    ]

    # 計算起始 y：從頂部排列
    total_h = len(steps) * (BH + GAP) - GAP
    start_y = 17.5 - total_h  # 頂部留 0.5 單位
    y_positions = [start_y + (len(steps)-1-i) * (BH + GAP) for i in range(len(steps))]

    for i, (num, title, params) in enumerate(steps):
        y = y_positions[i]
        sbox(y, num, title, params)
        if i < len(steps) - 1:
            darr(y)

    # 分支：合格 → 蝕刻，不合格 → 重工
    insp_y = y_positions[-1]  # ⑨ 檢測的 y
    branch_y = insp_y - GAP

    # 菱形判斷框
    cx, cy = 5.0, branch_y - 0.35
    pts = [(cx, cy+0.38), (cx+1.1, cy), (cx, cy-0.38), (cx-1.1, cy)]
    ax.add_patch(plt.Polygon(pts, fc='#F5F5F5', ec=EC, lw=1.3, zorder=2))
    ax.text(cx, cy, 'CD 合格？', ha='center', va='center',
            fontsize=9, fontweight='bold', color=TC, zorder=3)

    # 箭頭從 ⑨ 到菱形
    ax.annotate('', xy=(cx, cy+0.38), xytext=(cx, insp_y),
                arrowprops=dict(arrowstyle='->', color='#333333', lw=1.5,
                                mutation_scale=14), zorder=1)

    # 合格 → 往下繼續
    pass_y = cy - 0.38 - GAP - BH
    sbox(pass_y, '⑩', '蝕刻 / 離子佈植  Etching / Implant',
         '去除未受光阻保護之材料層，形成元件結構')
    ax.annotate('', xy=(cx, pass_y+BH), xytext=(cx, cy-0.38),
                arrowprops=dict(arrowstyle='->', color='#333333', lw=1.5,
                                mutation_scale=14), zorder=1)
    ax.text(cx+0.12, cy-0.38-GAP/2, '合格', fontsize=8.5,
            color='#1A1A1A', va='center')

    strip_y = pass_y - GAP - BH
    sbox(strip_y, '11', '光阻去除  Strip / Ash',
         '乾式灰化或濕式剝離，去除殘餘光阻')
    ax.annotate('', xy=(cx, strip_y+BH), xytext=(cx, pass_y),
                arrowprops=dict(arrowstyle='->', color='#333333', lw=1.5,
                                mutation_scale=14), zorder=1)

    next_y = strip_y - GAP - BH
    sbox(next_y, '12', '下一道製程  Next Layer', '', dark=True)
    ax.annotate('', xy=(cx, next_y+BH), xytext=(cx, strip_y),
                arrowprops=dict(arrowstyle='->', color='#333333', lw=1.5,
                                mutation_scale=14), zorder=1)

    # 不合格 → 重工 (向右彎回到 ③)
    rework_x = BX + BW + 0.35
    ax.annotate('', xy=(rework_x, cy), xytext=(cx+1.1, cy),
                arrowprops=dict(arrowstyle='->', color='#888888', lw=1.3,
                                mutation_scale=12), zorder=1)
    ax.text(cx+1.15, cy+0.08, '不合格', fontsize=8, color='#555555')
    coat_y = y_positions[2]  # ③ 光阻塗佈
    ax.annotate('', xy=(rework_x, coat_y+BH/2), xytext=(rework_x, cy),
                arrowprops=dict(arrowstyle='->', color='#888888', lw=1.3,
                                mutation_scale=12,
                                connectionstyle='arc3,rad=0'), zorder=1)
    ax.annotate('', xy=(BX+BW, coat_y+BH/2), xytext=(rework_x, coat_y+BH/2),
                arrowprops=dict(arrowstyle='->', color='#888888', lw=1.3,
                                mutation_scale=12), zorder=1)
    ax.text(rework_x+0.08, (cy+coat_y+BH/2)/2, '重工\n(Rework)',
            fontsize=8, color='#555555', va='center')

    ax.set_title('圖 6　DUV 微影製程步驟流程圖', fontsize=13,
                 fontweight='bold', pad=10, color='#1A1A1A')
    fig.tight_layout(pad=0.5)
    return fig_buf(fig)

# ══════════════════════════════════════════════════════════════════════════════
# 圖 4：曝光模擬控制流程圖（對應實際實作）
# ══════════════════════════════════════════════════════════════════════════════
def fig_adaptive_flow():
    """自適應教學模式判斷流程圖（獨立圖，用於 2.2 節）"""
    from matplotlib.patches import FancyBboxPatch
    fig, ax = plt.subplots(figsize=(12, 17))
    fig.patch.set_facecolor('white')
    ax.set_facecolor('white')
    ax.set_xlim(0, 12)
    ax.set_ylim(0, 17)
    ax.axis('off')

    EC = '#2C2C2C'; TC = '#1A1A1A'; CX = 5.0
    BX, BW = 1.8, 6.4

    def cbox(x, y, w, h, title, params='', dark=False):
        fc = '#2C2C2C' if dark else 'white'
        tc = 'white' if dark else TC
        ax.add_patch(FancyBboxPatch((x, y), w, h,
                     boxstyle='square,pad=0', lw=1.3,
                     edgecolor=EC, facecolor=fc, zorder=2))
        ax.text(x+w/2, y+h*(0.65 if params else 0.5), title,
                ha='center', va='center', fontsize=9.5,
                fontweight='bold', color=tc, zorder=3, multialignment='center')
        if params:
            ax.text(x+w/2, y+h*0.25, params,
                    ha='center', va='center', fontsize=8,
                    color='#444' if not dark else '#BBB',
                    zorder=3, style='italic', multialignment='center')

    def diam(cx, cy, hw, hh, text):
        pts = [(cx, cy+hh), (cx+hw, cy), (cx, cy-hh), (cx-hw, cy)]
        ax.add_patch(plt.Polygon(pts, fc='#F5F5F5', ec=EC, lw=1.3, zorder=2))
        ax.text(cx, cy, text, ha='center', va='center',
                fontsize=8.5, fontweight='bold', color=TC, zorder=3,
                multialignment='center')

    def darr(y1, y2, x=None, label='', lc='#333333'):
        xp = x if x is not None else CX
        ax.annotate('', xy=(xp, y2), xytext=(xp, y1),
                    arrowprops=dict(arrowstyle='->', color=lc,
                                   lw=1.6, mutation_scale=14), zorder=1)
        if label:
            ax.text(xp+0.14, (y1+y2)/2, label,
                    fontsize=8.5, fontweight='bold', color=lc, va='center')

    def harr(x1, y, x2, label='', lc='#555555'):
        ax.annotate('', xy=(x2, y), xytext=(x1, y),
                    arrowprops=dict(arrowstyle='->', color=lc,
                                   lw=1.3, mutation_scale=12), zorder=1)
        if label:
            ax.text((x1+x2)/2, y+0.1, label, fontsize=8.5,
                    fontweight='bold', color=lc, ha='center')

    def mode_box(x, y, w, h, title, desc, shade='#F0F0F0'):
        ax.add_patch(FancyBboxPatch((x, y), w, h,
                     boxstyle='square,pad=0', lw=1.2,
                     edgecolor='#444', facecolor=shade, zorder=2))
        ax.text(x+w/2, y+h*0.68, title, ha='center', va='center',
                fontsize=9, fontweight='bold', color='#1A1A1A', zorder=3)
        ax.text(x+w/2, y+h*0.28, desc, ha='center', va='center',
                fontsize=7.5, color='#444', zorder=3, style='italic')

    # ── 1. 訓練情境啟動 ──────────────────────────────────────────────
    cbox(BX, 15.5, BW, 1.0,
         '訓練情境啟動  /api/start',
         'SECOM 數位孿生 590 感測器  ·  故障情境隨機選定  ·  AI 導師初始化',
         dark=True)
    darr(15.5, 14.85)

    # ── 2. 使用者操作 3D 環境 ─────────────────────────────────────────
    cbox(BX, 14.15, BW, 0.65,
         '使用者操作 3D 環境  (viewer.html)',
         'WASD 移動  ·  [E] 互動 Mesh  ·  [C] 開啟 Chat  ·  HMI 面板查看')
    darr(14.15, 13.5)

    # ── 3. 使用者回答 AI 問題 ─────────────────────────────────────────
    cbox(BX, 12.8, BW, 0.65,
         '使用者回答 AI 問題  /api/chat',
         'Qwen LLM 評估回答內容  ·  AI 評分 0–10 分')
    darr(12.8, 12.15)

    # ── 4. 故障診斷操作判斷 ──────────────────────────────────────────
    dcy1 = 11.67
    diam(CX, dcy1, 1.5, 0.45, '診斷正確？')
    darr(12.15, dcy1+0.45)

    # YES → 右側：記錄正確
    rbox_x = BX + BW + 0.2
    cbox(rbox_x, 11.27, 2.2, 0.8,
         '診斷正確\n記錄得分', '', dark=False)
    harr(CX+1.5, dcy1, rbox_x, 'YES', lc='#1A7340')

    # NO → 繼續
    darr(dcy1-0.45, 10.85, label='NO', lc='#C0392B')

    # ── 5. 連續答對 ≥ 3 → 升級？ ────────────────────────────────────
    dcy2 = 10.37
    diam(CX, dcy2, 1.5, 0.45, '連續答對\n≥ 3 次？')
    darr(10.85, dcy2+0.45)

    # YES → 挑戰模式
    mode_box(rbox_x, 9.97, 2.3, 0.78,
             '挑戰模式', '複雜整合題  ·  無提示  ·  簡潔反饋', shade='#E8F4E8')
    harr(CX+1.5, dcy2, rbox_x, 'YES', lc='#1A7340')

    # NO → 往下
    darr(dcy2-0.45, 9.55, label='NO', lc='#C0392B')

    # ── 6. 連續答錯 ≥ 3 → 降級？ ────────────────────────────────────
    dcy3 = 9.07
    diam(CX, dcy3, 1.5, 0.45, '連續答錯\n≥ 3 次？')
    darr(9.55, dcy3+0.45)

    mode_box(rbox_x, 8.67, 2.3, 0.78,
             '標準模式', '適中題目  ·  少量提示  ·  平衡反饋', shade='#EAF2FB')
    harr(CX+1.5, dcy3, rbox_x, 'NO', lc='#1A7340')

    darr(dcy3-0.45, 8.25, label='YES', lc='#C0392B')

    # ── 7. 再連續答錯 ≥ 3？ ──────────────────────────────────────────
    dcy4 = 7.77
    diam(CX, dcy4, 1.5, 0.45, '再連續\n答錯 ≥ 3？')
    darr(8.25, dcy4+0.45)

    mode_box(rbox_x, 7.37, 2.3, 0.78,
             '鷹架模式', '基礎題目  ·  頻繁提示  ·  詳細反饋', shade='#FEF9E7')
    harr(CX+1.5, dcy4, rbox_x, 'NO', lc='#1A7340')

    darr(dcy4-0.45, 7.0, label='YES', lc='#C0392B')

    # 補救模式（直接在主流）
    mode_box(BX, 6.3, BW, 0.65,
             '補救模式  Remedial Mode',
             '最簡單題目  ·  大量提示  ·  超詳細反饋  ·  適合初學者',
             shade='#FDEDEC')
    darr(6.3, 5.65)

    # ── 8. 教學模式合流 → 維修 SOP ───────────────────────────────────
    # 從右側四個 mode_box 連回主流
    for my in [10.36, 9.06, 7.76]:
        ax.annotate('', xy=(BX+BW, 6.62),
                    xytext=(rbox_x+2.3, my),
                    arrowprops=dict(arrowstyle='->', color='#AAAAAA', lw=1.0,
                                   connectionstyle='arc3,rad=-0.25'), zorder=0)

    cbox(BX, 4.95, BW, 0.65,
         '零件互動  自主操作評估  /api/action',
         '受訓者自選零件與操作  ·  AI 即時比對 SOP 序列  ·  求助學長 /api/hint（-5 分）')
    darr(4.95, 4.3)

    # ── 9. ScoringSystem ─────────────────────────────────────────────
    cbox(BX, 3.3, BW, 0.95,
         'ScoringSystem 評分',
         '診斷準確性 30%  ·  操作正確性 40%  ·  安全合規 30%  ·  時間效率加分')
    darr(3.3, 2.65)

    # ── 10. 完成判斷 ─────────────────────────────────────────────────
    dcy5 = 2.17
    diam(CX, dcy5, 1.5, 0.45, '總分\n≥ 80 分？')
    darr(2.65, dcy5+0.45)

    darr(dcy5-0.45, 1.5, label='NO  重新訓練', lc='#C0392B')
    cbox(BX, 0.5, BW, 0.95,
         '訓練完成  能力認證  &  進階推薦',
         '等級評定 A–F  ·  Competency Assessment  ·  智慧推薦下一訓練模組',
         dark=True)
    harr(CX+1.5, dcy5, BX+BW+0.1, 'YES  完成', lc='#1A7340')
    ax.annotate('', xy=(BX+BW, 0.97),
                xytext=(BX+BW+0.1, dcy5),
                arrowprops=dict(arrowstyle='->', color='#1A7340', lw=1.3,
                                connectionstyle='arc3,rad=-0.3'), zorder=1)

    ax.set_title('圖 2　自適應教學模式判斷流程圖', fontsize=13,
                 fontweight='bold', pad=10, color='#1A1A1A')
    fig.tight_layout(pad=0.5)
    return fig_buf(fig)

# ══════════════════════════════════════════════════════════════════════════════
# 圖 5：DUV 曝光站光學路徑圖
# ══════════════════════════════════════════════════════════════════════════════
def fig_optical_path():
    """DUV 曝光站光學路徑圖（僅含本專案實作元件）"""
    from matplotlib.patches import FancyBboxPatch
    fig, ax = plt.subplots(figsize=(8, 11))
    fig.patch.set_facecolor('white')
    ax.set_facecolor('white')
    ax.set_xlim(0, 8)
    ax.set_ylim(0, 11)
    ax.axis('off')

    CX = 3.5
    BX, BW = 1.0, 5.0
    EC = '#2C2C2C'; TC = '#1A1A1A'

    def obox(y, h, title, params='', dark=False, ec_col=None, fc_col=None):
        fc = fc_col if fc_col else ('#2C2C2C' if dark else 'white')
        tc = 'white' if dark else TC
        ec = ec_col if ec_col else EC
        lw = 1.8 if ec_col else 1.3
        ax.add_patch(FancyBboxPatch((BX, y), BW, h,
                     boxstyle='square,pad=0', lw=lw,
                     edgecolor=ec, facecolor=fc, zorder=2))
        ax.text(CX, y+h*(0.65 if params else 0.5), title,
                ha='center', va='center', fontsize=10,
                fontweight='bold', color=tc, zorder=3, multialignment='center')
        if params:
            ax.text(CX, y+h*0.26, params, ha='center', va='center',
                    fontsize=8.5, color=tc if dark else '#555',
                    zorder=3, style='italic', multialignment='center')

    def darr(y1, y2):
        ax.annotate('', xy=(CX, y2), xytext=(CX, y1),
                    arrowprops=dict(arrowstyle='->', color='#333333',
                                   lw=1.8, mutation_scale=16), zorder=3)

    # 光束通道（淡黃帶）
    ax.add_patch(FancyBboxPatch((CX-0.32, 1.4), 0.64, 8.8,
                 boxstyle='square,pad=0', lw=0, facecolor='#FFFBF0', zorder=0))
    ax.text(CX+0.55, 5.8, '光束方向  ↓', fontsize=7.5,
            color='#BBA000', va='center', rotation=270, style='italic')

    # 1. KrF 雷射光源
    obox(10.0, 0.8, 'KrF 準分子雷射  Excimer Laser',
         '波長 248 nm  ·  光源穩定度監控', dark=True)
    darr(10.0, 9.25)

    # 2. 照明系統（IllumLens_0~5，σ 控制）
    obox(8.5, 0.7, '照明系統  Illumination System  (IllumLens × 6)',
         'σ 相干度設定  ·  照明均勻度校正',
         ec_col='#1A5276', fc_col='#EBF5FB')
    # σ 控制標籤
    ax.annotate('σ 控制', xy=(BX, 8.85), xytext=(0.05, 8.85),
                fontsize=8, color='#1A5276', fontweight='bold', va='center',
                arrowprops=dict(arrowstyle='->', color='#1A5276', lw=1.0))
    darr(8.5, 7.75)

    # 3. 光罩 + 光罩載台（Reticle_Mesh, Reticle_Stage_Mesh）
    obox(7.0, 0.7, '光罩  Reticle  +  光罩載台  Reticle Stage',
         '4× 縮小電路圖案  ·  光罩污染故障情境（contamination）',
         ec_col='#1B5E20', fc_col='#F1F8F1')
    ax.text(BX+BW+0.1, 7.35, '4×', fontsize=11,
            color='#1B5E20', fontweight='bold', va='center')
    darr(7.0, 6.25)

    # 4. 投影鏡組（Lens_0~9，NA 控制，lens_hotspot 故障）
    obox(5.5, 0.7, '投影鏡組  Projection Lens  (Lens × 10)',
         'NA 數值孔徑控制  ·  鏡片過熱故障情境（lens_hotspot）',
         ec_col='#7D3C00', fc_col='#FEF0E6')
    # NA 控制標籤
    ax.annotate('NA 控制', xy=(BX, 5.85), xytext=(0.05, 5.85),
                fontsize=8, color='#7D3C00', fontweight='bold', va='center',
                arrowprops=dict(arrowstyle='->', color='#7D3C00', lw=1.0))
    darr(5.5, 4.2)

    # 5. 晶圓台
    obox(3.5, 0.7, '晶圓台  Wafer Stage  (Step & Scan)',
         '步進掃描曝光  ·  晶圓台位置誤差故障情境（stage_error）')
    darr(3.5, 2.7)

    # 6. 晶圓光阻層
    obox(1.7, 0.9, '晶圓光阻層  Photoresist on Wafer',
         '化學放大光阻（CAR）  ·  潛像形成', dark=True)

    ax.set_title('圖 5　DUV 曝光站光學路徑圖', fontsize=13,
                 fontweight='bold', pad=10, color='#1A1A1A')
    fig.tight_layout(pad=0.5)
    return fig_buf(fig)

# ══════════════════════════════════════════════════════════════════════════════
# 圖 6：訓練互動操作 + 自適應教學整合流程圖
# ══════════════════════════════════════════════════════════════════════════════
def fig_exposure_optical():
    """訓練互動操作（HMI/故障注入）+ 自適應教學整合流程圖"""
    from matplotlib.patches import FancyBboxPatch
    fig, ax = plt.subplots(figsize=(12, 22))
    fig.patch.set_facecolor('white')
    ax.set_facecolor('white')
    ax.set_xlim(0, 12)
    ax.set_ylim(0, 22)
    ax.axis('off')

    EC = '#2C2C2C'; TC = '#1A1A1A'; CX = 5.0
    BX, BW = 1.8, 6.4

    def cbox(x, y, w, h, title, params='', dark=False):
        fc = '#2C2C2C' if dark else 'white'
        tc = 'white' if dark else TC
        ax.add_patch(FancyBboxPatch((x, y), w, h,
                     boxstyle='square,pad=0', lw=1.3,
                     edgecolor=EC, facecolor=fc, zorder=2))
        ax.text(x+w/2, y+h*(0.65 if params else 0.5), title,
                ha='center', va='center', fontsize=9.5,
                fontweight='bold', color=tc, zorder=3, multialignment='center')
        if params:
            ax.text(x+w/2, y+h*0.25, params,
                    ha='center', va='center', fontsize=8,
                    color='#444' if not dark else '#BBB',
                    zorder=3, style='italic', multialignment='center')

    def diam(cx, cy, hw, hh, text):
        pts = [(cx, cy+hh), (cx+hw, cy), (cx, cy-hh), (cx-hw, cy)]
        ax.add_patch(plt.Polygon(pts, fc='#F5F5F5', ec=EC, lw=1.3, zorder=2))
        ax.text(cx, cy, text, ha='center', va='center',
                fontsize=8.5, fontweight='bold', color=TC, zorder=3,
                multialignment='center')

    def darr(y1, y2, x=None, label='', lc='#333333'):
        xp = x if x is not None else CX
        ax.annotate('', xy=(xp, y2), xytext=(xp, y1),
                    arrowprops=dict(arrowstyle='->', color=lc,
                                   lw=1.6, mutation_scale=14), zorder=1)
        if label:
            ax.text(xp+0.14, (y1+y2)/2, label,
                    fontsize=8.5, fontweight='bold', color=lc, va='center')

    def harr(x1, y, x2, label='', lc='#555555'):
        ax.annotate('', xy=(x2, y), xytext=(x1, y),
                    arrowprops=dict(arrowstyle='->', color=lc,
                                   lw=1.3, mutation_scale=12), zorder=1)
        if label:
            ax.text((x1+x2)/2, y+0.1, label, fontsize=8.5,
                    fontweight='bold', color=lc, ha='center')

    def mode_box(x, y, w, h, title, desc, shade='#F0F0F0'):
        ax.add_patch(FancyBboxPatch((x, y), w, h,
                     boxstyle='square,pad=0', lw=1.2,
                     edgecolor='#444', facecolor=shade, zorder=2))
        ax.text(x+w/2, y+h*0.68, title, ha='center', va='center',
                fontsize=9, fontweight='bold', color='#1A1A1A', zorder=3)
        ax.text(x+w/2, y+h*0.28, desc, ha='center', va='center',
                fontsize=7.5, color='#444', zorder=3, style='italic')

    # ── 上半段：訓練互動操作層 ────────────────────────────────────────
    cbox(BX, 21.0, BW, 0.9,
         '訓練情境啟動  /api/start',
         'SECOM 數位孿生 590 感測器  ·  故障情境隨機選定  ·  AI 導師初始化',
         dark=True)
    darr(21.0, 20.35)

    cbox(BX, 19.7, BW, 0.6,
         '使用者操作 3D 環境  (viewer.html)',
         'WASD 移動  ·  [E] 互動 Mesh  ·  [C] Chat  ·  HMI 面板查看')
    darr(19.7, 19.05)

    cbox(BX, 18.0, BW, 0.8,
         '故障情境注入設定  /api/fault',
         '故障類型：lens_hotspot · contamination · stage_error · dose_drift · focus_drift')
    darr(18.0, 17.35)

    cbox(BX, 16.7, BW, 0.6,
         'HMI 面板即時更新',
         'SECOM 雜訊模型觸發  ·  感測器異常值顯示  ·  警報啟動')
    darr(16.7, 16.05)

    cbox(BX, 15.4, BW, 0.6,
         'HMI 即時監控',
         'cd_mean · pass/fail 狀態  ·  590 感測器讀值  ·  警報通知')
    darr(15.4, 14.75)

    cbox(BX, 14.1, BW, 0.6,
         '使用者回答 AI 問題  /api/chat',
         'Qwen LLM 評估回答內容  ·  AI 評分 0–10 分')
    darr(14.1, 13.52)

    # 分層虛線標記
    ax.plot([BX-0.5, BX+BW+2.6], [13.8, 13.8],
            color='#BBBBBB', lw=0.9, ls='--', zorder=0)
    ax.text(BX-0.45, 13.85, '▲ 操作互動', fontsize=7.5, color='#999', va='bottom')
    ax.text(BX-0.45, 13.75, '▼ 自適應判斷', fontsize=7.5, color='#999', va='top')

    # ── 下半段：自適應教學判斷層 ──────────────────────────────────────
    rbox_x = BX + BW + 0.2

    dcy1 = 13.07
    diam(CX, dcy1, 1.5, 0.45, '診斷正確？')
    darr(13.52, dcy1+0.45)
    cbox(rbox_x, 12.67, 2.2, 0.78, '診斷正確\n記錄得分', '', dark=False)
    harr(CX+1.5, dcy1, rbox_x, 'YES', lc='#1A7340')
    darr(dcy1-0.45, 12.17, label='NO', lc='#C0392B')

    dcy2 = 11.72
    diam(CX, dcy2, 1.5, 0.45, '連續答對\n≥ 3 次？')
    darr(12.17, dcy2+0.45)
    mode_box(rbox_x, 11.32, 2.3, 0.78, '挑戰模式', '複雜整合題  ·  無提示  ·  簡潔反饋', shade='#E8F4E8')
    harr(CX+1.5, dcy2, rbox_x, 'YES', lc='#1A7340')
    darr(dcy2-0.45, 10.82, label='NO', lc='#C0392B')

    dcy3 = 10.37
    diam(CX, dcy3, 1.5, 0.45, '連續答錯\n≥ 3 次？')
    darr(10.82, dcy3+0.45)
    mode_box(rbox_x, 9.97, 2.3, 0.78, '標準模式', '適中題目  ·  少量提示  ·  平衡反饋', shade='#EAF2FB')
    harr(CX+1.5, dcy3, rbox_x, 'NO', lc='#1A7340')
    darr(dcy3-0.45, 9.47, label='YES', lc='#C0392B')

    dcy4 = 9.02
    diam(CX, dcy4, 1.5, 0.45, '再連續\n答錯 ≥ 3？')
    darr(9.47, dcy4+0.45)
    mode_box(rbox_x, 8.62, 2.3, 0.78, '鷹架模式', '基礎題目  ·  頻繁提示  ·  詳細反饋', shade='#FEF9E7')
    harr(CX+1.5, dcy4, rbox_x, 'NO', lc='#1A7340')
    darr(dcy4-0.45, 8.12, label='YES', lc='#C0392B')

    mode_box(BX, 7.50, BW, 0.62, '補救模式  Remedial Mode',
             '最簡單題目  ·  大量提示  ·  超詳細反饋  ·  適合初學者', shade='#FDEDEC')
    darr(7.50, 6.80)

    # 右側教學模式 → SOP 彙流
    for my in [11.71, 10.36, 9.01]:
        ax.annotate('', xy=(BX+BW, 7.09),
                    xytext=(rbox_x+2.3, my),
                    arrowprops=dict(arrowstyle='->', color='#AAAAAA', lw=1.0,
                                   connectionstyle='arc3,rad=-0.25'), zorder=0)

    cbox(BX, 6.15, BW, 0.65,
         '零件互動  自主操作評估  /api/action',
         '受訓者自選零件與操作  ·  AI 即時比對 SOP 序列  ·  求助學長 /api/hint（-5 分）')
    darr(6.15, 5.45)

    cbox(BX, 4.45, BW, 0.95,
         'ScoringSystem 評分',
         '診斷準確性 30%  ·  操作正確性 40%  ·  安全合規 30%  ·  時間效率加分')
    darr(4.45, 3.80)

    dcy5 = 3.35
    diam(CX, dcy5, 1.5, 0.45, '總分\n≥ 80 分？')
    darr(3.80, dcy5+0.45)
    darr(dcy5-0.45, 1.40, label='NO  重新訓練', lc='#C0392B')

    cbox(BX, 0.5, BW, 0.9,
         '訓練完成  能力認證  &  進階推薦',
         '等級評定 A–F  ·  Competency Assessment  ·  智慧推薦下一訓練模組',
         dark=True)
    harr(CX+1.5, dcy5, BX+BW+0.1, 'YES  完成', lc='#1A7340')
    ax.annotate('', xy=(BX+BW, 0.95),
                xytext=(BX+BW+0.1, dcy5),
                arrowprops=dict(arrowstyle='->', color='#1A7340', lw=1.3,
                                connectionstyle='arc3,rad=-0.3'), zorder=1)

    ax.set_title('圖 6　訓練互動操作與自適應教學整合流程圖', fontsize=13,
                 fontweight='bold', pad=10, color='#1A1A1A')
    fig.tight_layout(pad=0.5)
    return fig_buf(fig)

# ══════════════════════════════════════════════════════════════════════════════
# 圖（舊）：資料流程圖
# ══════════════════════════════════════════════════════════════════════════════
def fig_dataflow():
    fig, ax = plt.subplots(figsize=(13,5.5))
    fig.patch.set_facecolor('#F8F9FA'); ax.set_facecolor('#F8F9FA')
    ax.set_xlim(0,13); ax.set_ylim(0,5.5); ax.axis('off')

    # Frontend
    box(ax,0.15,1.0,2.5,3.5,'前端瀏覽器',
        'Three.js 3D 場景\nPointerLock WASD\nHMI 虛擬面板\nChat 對話框\nE鍵互動',
        fc='#1E3A5F',fs=8.8)
    # Arrow →
    arr(ax,2.65,2.75,3.4,2.75,'POST /api/exposure\nPOST /api/chat\nGET /api/hmi',
        c=BLUE,lw=2,loff=(-0.05,0.08))
    # API
    box(ax,3.4,1.0,2.5,3.5,'REST API\nPort 8765',
        '/api/start\n/api/chat\n/api/exposure\n/api/hmi\n/api/wafer_map\n/api/fault',
        fc='#1B4F20',fs=8.8)
    arr(ax,5.9,2.75,6.65,2.75,'路由分發',c=GREEN,lw=2)
    # Physics
    box(ax,6.65,3.2,2.6,1.3,'LensHeatingEngine\n物理引擎',
        '雙指數熱模型\nBossung → CD\nZernike → Overlay',
        fc=RED,fs=8.5)
    box(ax,6.65,1.55,2.6,1.3,'SecomNoiseModel\nUCI SECOM',
        '590 特徵 PCA 分析\nhealth_score → noise σ\n5 種故障模型',
        fc=PURPLE,fs=8.5)
    # noise sigma
    ax.annotate('',xy=(7.95,3.2),xytext=(7.95,2.85),
                arrowprops=dict(arrowstyle='->',color=PURPLE,lw=1.8))
    ax.text(8.05,3.0,'noise σ',fontsize=7.5,color=PURPLE)
    arr(ax,9.25,3.85,10.1,3.1,c=RED,lw=1.8)
    arr(ax,9.25,2.2,10.1,2.9,c=PURPLE,lw=1.8)
    # Response
    box(ax,10.1,2.0,2.7,2.0,'JSON 回傳結果',
        'cd_mean · cd_3σ\noverlay_x/y_3σ\ncdu_map[13×13]\nzernike{Z4,Z7,Z9}\nyield_prob · pass_fail',
        fc=TEAL,fs=8.5)
    arr(ax,10.1,3.0,5.9,3.2,'JSON Response',c=ORANGE,lw=2,loff=(0.1,0.12))
    arr(ax,3.4,2.75,2.65,2.75,'ai_msg + data',c=ORANGE,lw=2)
    # Render output
    box(ax,0.15,0.1,2.5,0.8,'渲染輸出',
        'CDU Map · 趨勢圖 · 鏡片溫度 · Overlay 向量',
        fc='#0D5C63',fs=8)
    arr(ax,1.4,1.0,1.4,0.9,c=TEAL,lw=1.8)

    ax.set_title('圖 5　系統資料流程圖', fontsize=13,
                 fontweight='bold', pad=8, color=NAVY)
    fig.tight_layout(pad=0.4)
    return fig_buf(fig)

# ══════════════════════════════════════════════════════════════════════════════
# Word 文件建構
# ══════════════════════════════════════════════════════════════════════════════
doc = Document()
sec = doc.sections[0]
sec.page_width   = Cm(21.0); sec.page_height  = Cm(29.7)
sec.left_margin  = Cm(3.0);  sec.right_margin  = Cm(2.5)
sec.top_margin   = Cm(2.5);  sec.bottom_margin = Cm(2.5)

def set_rfont(r, size=13, bold=None, italic=None, color=None):
    """中英混排字型：中文=標楷體，英文=Times New Roman"""
    r.font.name = 'Times New Roman'
    r.font.size = Pt(size)
    if bold is not None:   r.font.bold   = bold
    if italic is not None: r.font.italic = italic
    if color is not None:  r.font.color.rgb = color
    rPr = r._r.get_or_add_rPr()
    for existing in rPr.findall(qn('w:rFonts')):
        rPr.remove(existing)
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'),    'Times New Roman')
    rFonts.set(qn('w:hAnsi'),    'Times New Roman')
    rFonts.set(qn('w:eastAsia'), '標楷體')
    rFonts.set(qn('w:cs'),       '標楷體')
    rPr.insert(0, rFonts)

def h(text, lv=1):
    hp = doc.add_heading(text, level=lv)
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for r in hp.runs:
        r.font.name = '標楷體'
        r.font.color.rgb = RGBColor(0x1A,0x3A,0x5C)
    return hp

def para(text, indent=True, justify=True):
    p = doc.add_paragraph(text)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY if justify else WD_ALIGN_PARAGRAPH.LEFT
    if indent:
        p.paragraph_format.first_line_indent = Cm(0.75)
    for r in p.runs:
        set_rfont(r, 13)
    return p

def caption(text):
    p = doc.add_paragraph(text)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in p.runs:
        set_rfont(r, 11, italic=True, color=RGBColor(0x44,0x44,0x44))

def add_toc():
    """插入 Word 可更新的自動目錄（按 Ctrl+A → F9 更新頁碼）"""
    p = doc.add_paragraph()
    run = p.add_run()
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'begin')
    run._r.append(fldChar)
    instr = OxmlElement('w:instrText')
    instr.set(qn('xml:space'), 'preserve')
    instr.text = ' TOC \\o "1-3" \\h \\z \\u '
    run._r.append(instr)
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    run._r.append(fldChar2)
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar3)

def fig(buf, w=6.0, cap=''):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(buf, width=Inches(w))
    if cap: caption(cap)
    doc.add_paragraph()

SS_DIR = r'c:\Users\user\Desktop\在職碩\OneDrive - 長庚大學\長庚碩班\論文\screenshots'
def ss(filename, w=5.5, cap=''):
    """插入截圖（從 screenshots/ 資料夾）"""
    import os
    path = os.path.join(SS_DIR, filename)
    if not os.path.exists(path):
        return  # 截圖尚未備妥，略過
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(path, width=Inches(w))
    if cap: caption(cap)
    doc.add_paragraph()

def bullet(text, lv=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Cm(0.5+lv*0.5)
    r = p.add_run(text)
    set_rfont(r, 13)

def eq(text, num=''):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.font.name = 'Times New Roman'; r.font.size = Pt(13); r.font.bold = True
    r.font.color.rgb = RGBColor(0x1A,0x3A,0x5C)
    if num:
        p2 = doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        nr = p2.add_run(num)
        set_rfont(nr, 11, italic=True)

# ── 封面 ──────────────────────────────────────────────────────────────────────
tp = doc.add_paragraph()
tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
tr = tp.add_run('半導體微影設備虛擬訓練系統')
tr.font.name='標楷體'; tr.font.size=Pt(22); tr.font.bold=True
tr.font.color.rgb = RGBColor(0x1A,0x3A,0x5C)

tp2 = doc.add_paragraph()
tp2.alignment = WD_ALIGN_PARAGRAPH.CENTER
tr2 = tp2.add_run('系統架構與設計說明')
tr2.font.name='標楷體'; tr2.font.size=Pt(16)
tr2.font.color.rgb = RGBColor(0x1A,0x3A,0x5C)

doc.add_paragraph()
for line in ['長庚大學　碩士論文','2024 – 2026']:
    ap = doc.add_paragraph(line)
    ap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in ap.runs:
        r.font.name='標楷體'; r.font.size=Pt(13)
        r.font.color.rgb = RGBColor(0x55,0x55,0x55)

doc.add_paragraph()
hrp = doc.add_paragraph('─' * 72)
hrp.alignment = WD_ALIGN_PARAGRAPH.CENTER
for r in hrp.runs: r.font.size=Pt(9); r.font.color.rgb=RGBColor(0x1A,0x3A,0x5C)

# ── 摘要 ──────────────────────────────────────────────────────────────────────
doc.add_paragraph()
h('摘要', 1)
para('本研究開發一套以沉浸式三維互動環境為核心之半導體微影設備虛擬訓練系統，'
     '以 ASML DUV（深紫外光，248 nm）微影機台為目標設備。'
     '系統整合 UCI SECOM 真實製程資料集（1,567 晶圓樣本、590 個感測器特徵）'
     '建立設備數位孿生（Lithography Digital Twin），提供感測器即時模擬與製程狀態追蹤。'
     'AI 導師系統採用本地端 Qwen 大型語言模型，結合 Diagnostic Agent、Operation Agent '
     '與 Adaptive Teaching 多代理架構，依學員 AI 評分（0–10 分）自動切換'
     '挑戰、標準、鷹架、補救四種教學模式，提供個人化蘇格拉底式引導。'
     '訓練課程採整合式設計，理論問答與 3D 實機操作於同一情境中同步進行，'
     'AI 導師依據受訓者輸入內容動態切換理論引導或實作指引，無分離的學科測驗門檻。'
     '實機操作採「自主評估」設計：受訓者自行判斷應對哪個零件執行何種操作，'
     'AI 導師即時評估操作正確性並依答對/答錯連續次數自動切換教學模式；'
     '設有「求助學長」提示機制（-5 分扣點）支援不同能力的學員。'
     '搭配三維度評分系統（診斷準確性 30%、操作正確性 40%、安全合規 30%）進行量化評量，'
     '最終以 A–F 五級制認證學員能力。')
doc.add_paragraph()
kp = doc.add_paragraph()
kr = kp.add_run('關鍵詞：')
set_rfont(kr, 13, bold=True)
kp2 = kp.add_run('半導體微影、DUV 虛擬訓練系統、數位孿生、SECOM 資料集、'
                 'Qwen LLM、自適應教學、Three.js、第一人稱訓練環境')
set_rfont(kp2, 13, italic=True)

doc.add_page_break()

# ── 目錄 ───────────────────────────────────────────────────────────────────────
h('目錄', 1)
add_toc()
doc.add_page_break()

# ── 一、緒論 ──────────────────────────────────────────────────────────────────
h('一、緒論', 1)
para('半導體微影製程是積體電路（IC）製造中技術難度最高、成本最昂貴的關鍵步驟之一。'
     '微影製程成本約佔晶圓廠整體製造成本的 30–35%，而 ASML DUV 微影設備單台採購成本'
     '即逾數千萬美元，設備停機一小時的機會成本可達數萬至十數萬美元。'
     '操作此類設備需具備深厚的光學物理、製程控制與故障診斷知識，'
     '傳統訓練方式高度依賴實體設備機時與資深工程師全程現場指導：'
     '新進操作員通常需要 2–4 週的實機訓練才具備獨立上機能力，'
     '期間若發生劑量設定錯誤、光罩裝卸失誤或錯誤處理系統異常，'
     '輕則觸發設備聯鎖（interlock）停機，重則導致批量晶圓報廢，損失難以估計。')
para('本研究提出一套以「設備數位孿生」（Lithography Digital Twin）為核心之沉浸式虛擬訓練平台。'
     '系統透過第一人稱 3D 互動環境讓受訓者在無需接觸實體設備的情況下，'
     '自由探索機台結構、查看感測器即時讀值、操作虛擬 HMI（人機介面），'
     '並在 AI 導師引導下執行故障診斷與維修程序。'
     '系統以 UCI SECOM 真實製程資料集驅動數位孿生感測器模型，'
     '使訓練情境的感測器行為具備真實統計依據，'
     '同時整合物理模型（鏡片熱方程式、Bossung Curve）模擬製程退化現象，'
     '力求使受訓者在虛擬環境中獲得與實機相近的診斷與處置體驗。')
para('本系統之主要貢獻包括：'
     '（1）整合 UCI SECOM 真實資料集建立設備數位孿生與感測器即時模擬；'
     '（2）以 WebGL/Three.js 實作具互動 Mesh 的第一人稱 3D 訓練環境；'
     '（3）基於本地端 Qwen LLM 之多代理 AI 導師系統，含自適應四模式教學；'
     '（4）自主式故障操作評估設計，受訓者自行判斷操作順序，AI 即時評分回饋並設有求助提示機制（-5 分）；'
     '以及（5）三維度量化評分系統與 A–F 能力認證機制。')

# ── 二、目標設備介紹 ──────────────────────────────────────────────────────────
h('二、目標設備：ASML DUV 微影機台', 1)
para('本研究以 ASML TWINSCAN NXT 系列深紫外光（DUV）微影機台作為虛擬訓練之目標設備。'
     'ASML TWINSCAN 採雙平台掃描架構，一個平台進行晶圓量測對準、另一個同步曝光，'
     '以最大化產能並確保套刻精度。機台以 KrF 準分子雷射（波長 248 nm）為光源，'
     '透過精密光學系統將光罩（Reticle）上的電路圖案以 4:1 比例縮小投影至晶圓光阻層。'
     '整台設備由多個高精密子系統協同運作，任一子系統異常均可能導致良率下降，'
     '因此工程師需熟悉各模組功能與故障特徵，本系統即以此需求為訓練目標。')

h('2.1　機台主要結構模組', 2)
para('ASML DUV 機台由下列核心模組組成，如表 1 所示。本訓練系統的 3D 模型'
     '（asml_duv.glb）對各模組均設有可互動的 Mesh 群組，受訓者可在虛擬環境中'
     '直接接近並按 [E] 鍵查看模組狀態說明，對應關係亦列於表 1。')
doc.add_paragraph()
tbl_eq = doc.add_table(rows=1, cols=3)
tbl_eq.style = 'Table Grid'
for cell, t in zip(tbl_eq.rows[0].cells, ['模組名稱', '功能說明', '3D 模型 Mesh']):
    cell.text = t
    for r in cell.paragraphs[0].runs:
        set_rfont(r, 10, bold=True)
eq_rows = [
    ('KrF 準分子雷射光源\n（Excimer Laser）',
     '輸出波長 248 nm 之深紫外脈衝雷射，為曝光製程提供高能量光源',
     'Laser_Box / Laser_Out'),
    ('照明系統\n（Illumination System）',
     '均勻化雷射光強分布，透過 σ（部分相干度）設定控制照明條件，含 6 組照明鏡片',
     'Illum_Barrel\nIllumLens_0 ~ IllumLens_5'),
    ('光罩載台\n（Reticle Stage）',
     '承載石英光罩，在掃描曝光時執行高速往復定位，精度決定 CD 均一性',
     'Reticle_Stage'),
    ('光罩\n（Reticle / Photo-Mask）',
     '石英基板上鍍鉻電路圖案，4× 縮小比例，為圖案轉移的關鍵元件；光罩污染為常見故障',
     'Reticle'),
    ('投影鏡組\n（Projection Lens）',
     '由 10 組高精密鏡片元件組成，NA = 0.75，將光罩圖案以 4:1 縮小投影至晶圓；'
     '鏡片過熱（lens_hotspot）為訓練故障情境之一',
     'POB_Barrel\nLens_0 ~ Lens_9'),
    ('晶圓台\n（Wafer Stage）',
     '採 Step & Scan 步進掃描曝光，定位精度 < 1 nm；'
     '晶圓台位置誤差（stage_error）為訓練故障情境之一',
     'Wafer_Chuck / Stage_Base'),
]
for name, func, mesh in eq_rows:
    row = tbl_eq.add_row().cells
    row[0].text = name; row[1].text = func; row[2].text = mesh
    for c in row:
        for r in c.paragraphs[0].runs:
            set_rfont(r, 10)
caption('表 1　ASML DUV 機台主要模組與 3D 訓練模型對應關係')
doc.add_paragraph()

h('2.2　主要技術規格', 2)
para('表 2 列出本訓練系統所對應之 DUV 機台關鍵技術規格。'
     '這些參數直接影響微影製程品質，亦是訓練系統中感測器監控與故障診斷的核心依據。')
doc.add_paragraph()
tbl_spec = doc.add_table(rows=1, cols=2)
tbl_spec.style = 'Table Grid'
for cell, t in zip(tbl_spec.rows[0].cells, ['技術項目', '規格 / 說明']):
    cell.text = t
    for r in cell.paragraphs[0].runs:
        set_rfont(r, 10, bold=True)
specs = [
    ('雷射波長',        'KrF 248 nm 深紫外光'),
    ('數值孔徑（NA）',  '0.75（訓練系統可調）'),
    ('部分相干度（σ）', '0.3 ~ 0.9（訓練系統可調）'),
    ('理論解析度',      'R = kλ/NA，約 120 nm（k = 0.4）'),
    ('晶圓尺寸',        '300 mm'),
    ('曝光模式',        'Step & Scan 步進掃描'),
    ('SECOM 感測器數',  '590 個製程感測器特徵（UCI SECOM 資料集）'),
    ('訓練故障情境',    '5 種：鏡片熱點、光罩污染、\n晶圓台誤差、劑量漂移、焦距漂移'),
]
for item, val in specs:
    row = tbl_spec.add_row().cells
    row[0].text = item; row[1].text = val
    for c in row:
        for r in c.paragraphs[0].runs:
            set_rfont(r, 10)
caption('表 2　DUV 機台關鍵技術規格')
doc.add_paragraph()

# ── 三、微影製程流程 ──────────────────────────────────────────────────────────
h('三、微影製程流程', 1)
para('DUV 微影製程由晶圓基板準備至下一道製程共包含十二個主要步驟，'
     '圖 1 以學術標準流程圖呈現各步驟的先後順序及關鍵製程參數；'
     '圖 2 放大曝光站之詳細子流程與各量測監控項目；'
     '圖 3 呈現 DUV 曝光站光學路徑，依序經過照明系統（σ 控制）、'
     '光罩（4× 縮小圖案）、投影鏡組（NA 控制），最終投影至晶圓光阻層；'
     '圖 4 則整合訓練互動操作與自適應教學判斷流程，呈現故障注入至 AI 評分驅動之'
     '自適應模式選擇（挑戰/標準/鷹架/補救）、SOP 維修與評分認證的完整訓練互動邏輯。')
doc.add_paragraph()
fig(fig_process_steps(), w=5.2, cap='圖 1　DUV 微影製程步驟流程圖')
doc.add_paragraph()
fig(fig_process(), w=6.4, cap='圖 2　微影製程流程圖（DUV 曝光站詳細）')
doc.add_paragraph()
fig(fig_optical_path(), w=4.0, cap='圖 3　DUV 曝光站光學路徑圖（雷射 → 晶圓）')
doc.add_paragraph()
fig(fig_exposure_optical(), w=5.0, cap='圖 4　訓練互動操作與自適應教學整合流程圖')

para('各步驟說明如下：')
steps = [
    ('基板清洗（Wafer Clean）',   '去除晶圓表面顆粒與有機污染，確保後續薄膜附著性。'),
    ('底層薄膜沉積',              '沉積硬光罩（Hard Mask）或反射防止塗層（ARC），提升線寬均勻性。'),
    ('光阻塗佈（PR Coating）',    '旋塗底層抗反射塗層（BARC）與化學放大光阻（CAR），'
                                   '膜厚均勻性直接影響 CD 控制能力。'),
    ('軟烘烤（Soft Bake / PAB）', '驅除光阻中殘留溶劑，穩定光阻膜層結構。'),
    ('曝光（DUV 248nm Exposure）','以 DUV ArF 雷射透過光罩將電路圖案投影至晶圓表面，'
                                   '為整個製程中參數最多、技術最複雜的步驟。'),
    ('曝後烘烤（PEB）',           '活化光阻中的酸催化反應，擴散酸性物質以改善站立波效應。'),
    ('顯影（Develop）',           '以顯影液溶解曝光區域（正型）或未曝光區域（負型）光阻。'),
    ('檢測（Inspection / OCD）',  '以 CD-SEM 或光學臨界尺寸（OCD）量測確認線寬符合規格，'
                                   '不合格品送回重工，合格品繼續後段製程。'),
]
for name, desc in steps:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    br = p.add_run(name + '：')
    set_rfont(br, 13, bold=True)
    nr = p.add_run(desc)
    set_rfont(nr, 13)

# ── 四、系統架構 ──────────────────────────────────────────────────────────────
h('四、系統架構', 1)
para('系統採用六層式垂直架構設計，如圖 5 所示。架構由資料來源層向下流經 AI 導師系統、'
     '後端服務層、使用者介面層、訓練流程層，最終進入評分系統層，'
     '每一層具有明確定義的責任範疇並透過標準化介面與相鄰層溝通。')
doc.add_paragraph()
fig(fig_arch_paper(), w=6.0, cap='圖 5　系統整體架構圖')

h('4.1　資料來源層', 2)
para('資料來源層包含兩個核心元件：UCI SECOM 製程資料集與設備數位孿生（Lithography Digital Twin）。'
     'SECOM 資料集提供 1,567 筆真實晶圓量測記錄與 590 個感測器特徵，'
     '作為製程雜訊模型之統計基礎；數位孿生則模擬 590 個感測器的即時狀態，'
     '供 HMI 面板即時顯示與製程狀態追蹤使用。')

h('4.2　AI 導師系統層', 2)
para('AI 導師系統採用本地端 Qwen 大型語言模型，由三個核心模組組成。'
     '（1）AIScenarioMentor：實現蘇格拉底式對話引導，依場景脈絡產生差異化回應；'
     '（2）Diagnostic Agent / Operation Agent：負責故障根本原因分析、逐步操作引導與維修 SOP 管理；'
     '（3）Adaptive Teaching / Competency Assessment：依 AI 評分（0–10 分）'
     '動態切換挑戰（≥8.5）、標準（≥6.5）、鷹架（≥4.0）、補救（<4.0）四種教學模式，'
     '並持續記錄學習歷程。')
doc.add_paragraph()
fig(fig_adaptive_flow(), w=5.0, cap='圖 6　自適應教學模式判斷流程圖')

h('4.3　使用者介面層', 2)
para('使用者介面以單頁 WebGL 應用程式實作，透過本地 HTTP 伺服器提供服務。'
     '3D 場景採用 Three.js 渲染，搭配 PointerLock 第一人稱相機控制器（WASD 移動、[E] 鍵互動）。'
     '虛擬 HMI 螢幕以 HTML5 Canvas 2D API 繪製後作為貼圖材質即時更新，'
     '顯示感測器讀值、CDU Map 與 CD 趨勢圖；'
     'AI 對話面板（[C] 鍵開啟）提供即時問答；'
     '故障情境下，受訓者靠近零件按 [E] 鍵後，出現「🔍 檢查」與「⚙ 操作」分類按鈕，'
     '自行決定操作內容，由 /api/action 進行 AI 即時評分回饋。'
     '介面右上角顯示即時分數（滿分 100 分），並設有「求助學長」按鈕（扣 5 分）供需要提示的受訓者使用。')
ss('ss_hmi_panel.png', w=5.8, cap='圖 7　HMI 虛擬控制面板（感測器讀值、CDU Map 即時顯示）')
ss('ss_ai_chat.png',   w=2.5, cap='圖 8　AI 對話面板（[C] 鍵開啟）')

h('4.4　後端服務層', 2)
para('後端通訊以 Python 內建 http.server 模組架設於 8765 埠，'
     '所有端點均回傳 JSON 格式回應，共提供 4 個 GET 端點與 4 個 POST 端點，'
     '詳見表 3。')

doc.add_paragraph()
tbl = doc.add_table(rows=1, cols=3)
tbl.style = 'Table Grid'
for cell, t in zip(tbl.rows[0].cells, ['端點路徑','方法','功能說明']):
    cell.text = t
    for r in cell.paragraphs[0].runs:
        set_rfont(r, 10, bold=True)
apis = [
    ('/api/start',         'POST', '初始化故障情境；AI 導師開場說明'),
    ('/api/chat',          'POST', '送出使用者對話；Qwen LLM 產生 AI 回應'),
    ('/api/exposure',      'POST', '觸發物理模擬：dose, focus, NA, sigma → 完整結果'),
    ('/api/fault',         'POST', '注入或清除製程故障，影響 SECOM 雜訊模型'),
    ('/api/action',        'POST', '受訓者對零件執行操作；AI 評估正確性、更新分數、切換自適應模式'),
    ('/api/hint',          'POST', '求助學長：回傳當前步驟提示，扣除 5 分'),
    ('/api/hmi',           'GET',  '回傳 HMI 面板資料：590 感測器值、警報、SECOM 狀態'),
    ('/api/wafer_map',     'GET',  '回傳 CDU + Overlay 歷史資料（最近 30 片晶圓）'),
    ('/api/process_window','GET',  '回傳製程窗口：焦距範圍、DOF、劑量寬容度'),
    ('/api/secom_summary', 'GET',  '回傳 SECOM 特徵摘要與良率統計'),
]
for ep,m,d in apis:
    row = tbl.add_row().cells
    row[0].text=ep; row[1].text=m; row[2].text=d
    for c in row:
        for r in c.paragraphs[0].runs:
            set_rfont(r, 10)
caption('表 3　REST API 端點規格')
doc.add_paragraph()

h('4.5　3D 訓練環境實作成果', 2)
para('本系統以 Three.js WebGL 技術於瀏覽器端建構全沉浸式 3D 訓練環境，'
     '受訓者透過 PointerLock API 進入第一人稱視角，以 WASD 鍵移動、滑鼠旋轉視角、'
     '[E] 鍵與機台零件進行互動，[C] 鍵開啟 AI 對話面板。'
     '系統共建置 ASML DUV 微影機台之高精度 3D 模型（GLB 格式），'
     '涵蓋 11 組可互動零件群組，並支援 5 種製程故障情境的即時注入，'
     '包含鏡片熱點、光罩污染、載台誤差、劑量漂移及焦距漂移。'
     '虛擬 HMI 面板以 HTML5 Canvas 動態繪製，即時顯示 590 個感測器讀值、CDU Map 與 CD 趨勢圖，'
     '並於偵測到故障時觸發警報通知。'
     '故障發生後，受訓者需主動靠近相關零件並按 [E] 鍵開啟零件互動面板，'
     '面板依「🔍 檢查」與「⚙ 操作」兩類呈現可選動作；'
     '受訓者自行判斷正確操作順序，系統透過 /api/action 端點即時評估每一步操作的正確性，'
     '並依連續答對或答錯次數自動切換自適應教學模式，提供個人化 AI 引導回饋。')
ss('ss_3d_overview.png',  w=5.8, cap='圖 9　3D 訓練環境全景（第一人稱視角）')
ss('ss_3d_interact.png',  w=5.8, cap='圖 10　靠近零件後出現互動提示（[E] 鍵）')
ss('ss_fault_alarm.png',  w=4.0, cap='圖 11　故障注入後觸發警報通知')

h('4.6　3D 訓練模型零件介紹', 2)
para('本訓練系統的 3D 模型（asml_duv.glb）完整重現 ASML DUV 微影機台之主要子系統，'
     '共設有 11 個可互動零件群組。受訓者在第一人稱視角下靠近零件並按下 [E] 鍵，'
     '系統即顯示「檢查」與「操作」兩類可選動作，由受訓者自行判斷應執行哪一項動作，'
     'AI 導師即時評估操作是否符合當前 SOP 步驟並給予回饋。'
     '下圖為機台內部零件總覽，各子系統位置與名稱如圖所示。')
ss('元件介紹.jpg', w=5.8, cap='圖 12　3D 訓練環境可互動零件總覽（機台內部視角）')

# ── 零件詳細說明表 ──────────────────────────────────────────────────────────────
tbl_comp = doc.add_table(rows=1, cols=4)
tbl_comp.style = 'Table Grid'
for cell, t in zip(tbl_comp.rows[0].cells,
                   ['零件名稱', '功能說明', '代表 Mesh', '關聯故障']):
    cell.text = t
    for r in cell.paragraphs[0].runs:
        set_rfont(r, 10, bold=True)

comp_data = [
    ('雷射光源\nExcimer Laser',
     'KrF 準分子雷射（248 nm），提供曝光所需深紫外線光源。\n能量穩定性直接影響線寬（CD）均一性。',
     'Laser_Box\nLaser_Out',
     '劑量漂移\n(dose_drift)'),
    ('DUV 光束路徑\nBeam Path',
     '深紫外線光束從雷射到照明系統的傳輸路徑，\n需在氮氣環境中傳輸以避免 DUV 能量被空氣吸收衰減。',
     'Beam_H1/V1\nBeam_Spot',
     '—'),
    ('反射鏡\nMirror',
     '高精度光束折轉與準直鏡，鍍膜反射率 >99%，\n用於調整光束傳輸方向至照明系統入口。',
     'Mirror1~3',
     '—'),
    ('照明系統\nIllumination System',
     '將雷射光整形均勻化，確保光罩面均勻照射。\n包含 6 組照明鏡片（IllumLens_0~5）及快門控制。',
     'Illum_Barrel\nIllumLens_0~5',
     '焦距漂移\n(focus_drift)'),
    ('光罩載台\nReticle Stage',
     '承載石英光罩（Reticle），掃描曝光時執行高速往復定位。\n光罩移動精度直接決定圖案縮放比例與 CD 均一性。',
     'Reticle_Stage\nReticle',
     '光罩污染\n(contamination)'),
    ('投影鏡組\nProjection Optics',
     '由 10 組高精密鏡片（Lens_0~9）組成，NA = 0.75，\n將光罩圖案以 4:1 比例縮小投影至晶圓。\n鏡片溫升 0.01°C 即可能造成對焦偏移。',
     'POB_Barrel\nLens_0~9',
     '鏡片熱點\n(lens_hotspot)'),
    ('晶圓載台\nWafer Stage',
     'Step & Scan 精密定位平台，承載晶圓並執行步進掃描曝光。\n定位精度須達奈米等級，溫度或震動變化均影響對準。',
     'Wafer_Chuck\nStage_Base / Rail',
     '載台誤差\n(stage_error)'),
    ('晶圓傳送手臂\nRobot Arm',
     'SCARA 機械手臂，以伯努利原理無接觸方式夾持晶圓，\n負責從 FOUP 精確取出並放置到晶圓載台。',
     'Robot_Arm_Link\nRobot_Arm_Base',
     '—'),
    ('晶圓傳送埠\nFOUP Port',
     '前開式晶圓傳送盒（FOUP）介面，\n維持晶圓潔淨度並提供自動化裝卸，\n支援 25 片晶圓批量作業。',
     'FOUP_Port\nFOUP_Door',
     '—'),
    ('液浸系統\nImmersion Hood',
     '維持投影鏡組底部與晶圓間的超純水液浸層，\n提升等效 NA 以改善解析度，水溫需控制在 ±0.001°C。',
     'Immersion_Hood\nImmersion_Supply',
     '—'),
    ('控制系統\nControl Cabinet',
     '機台主控電腦與操作面板，即時監控並控制所有子系統，\n處理來自數千個感測器的即時數據；\n按 [E] 可開啟虛擬 HMI 控制面板。',
     'Cabinet_Main\nScreen / Keyboard',
     '—'),
]
for name, desc, mesh, fault in comp_data:
    row = tbl_comp.add_row().cells
    row[0].text = name
    row[1].text = desc
    row[2].text = mesh
    row[3].text = fault
    for c in row:
        for r in c.paragraphs[0].runs:
            set_rfont(r, 9.5)
caption('表 4　3D 訓練環境可互動零件群組詳細說明')
doc.add_paragraph()

# ── 五、SECOM 流程 ────────────────────────────────────────────────────────────
h('五、SECOM 資料驅動模擬流程', 1)
para('本系統之製程雜訊模型以 UCI SECOM 資料集為基礎，將真實晶圓製程的統計特性'
     '引入模擬環境，使數位孿生的感測器輸出具備真實統計依據。')

h('5.1　資料前處理', 2)
para('SECOM 資料集包含 1,567 筆晶圓量測記錄，每筆具有 590 個感測器特徵，'
     '標記為通過（1,463 筆）或失敗（104 筆）。首先以主成分分析（PCA）萃取'
     '與製程結果相關性最高的前 20 個特徵，並對各特徵的偏態分布進行建模，'
     '以準確重現實際製程雜訊的統計特性。')

h('5.2　健康分數與雜訊調變', 2)
para('依據 PCA 投影座標計算設備健康分數（health score），範圍為 0 至 1，'
     '數值越低代表設備狀態越差。健康分數直接調變 CD 預測的雜訊標準差：'
     '正常狀態下 σ = 1.5 nm，故障狀態下 σ = 3.0 nm，並依分數線性插值。'
     '此設計使模擬的製程雜訊具有真實統計依據，而非純粹合成分布。')

h('5.3　故障注入', 2)
para('系統支援五種製程故障類型的動態注入：dose_drift（劑量漂移）、'
     'focus_drift（焦距漂移）、lens_hotspot（鏡片熱點）、'
     'contamination（污染）及 stage_error（晶圓台位置誤差）。'
     '故障注入後，noise σ 即時升高，並反映至 CDU Map 及 Overlay 向量場的空間分布。')

# ── 六、SOP 動作評估系統 ──────────────────────────────────────────────────────
h('六、SOP 動作評估系統', 1)
para('本系統的實機操作訓練採「自主評估」設計，'
     '以預定義的 SOP 規則表（sop_definitions.py）作為操作正確性的判斷基準，'
     '搭配 Qwen LLM 生成自適應回饋文字。'
     '與 SECOM 資料集負責感測器統計特性不同，SOP 規則表負責定義「什麼操作步驟是正確的」。')

h('6.1　SOP 規則表結構', 2)
para('每種故障情境在 SOP_DEFINITIONS 中定義一組有序的步驟序列，'
     '每個步驟包含以下四個欄位：')
for name, desc in [
    ('valid_components', '此步驟應互動的零件群組（前端 label，精確比對）'),
    ('valid_actions',    '接受的動作關鍵詞清單（模糊比對，包含即算正確）'),
    ('hint_component',   '求助學長或答錯時給出的零件方向提示'),
    ('reason',           '此步驟的工程原因說明（補救模式下提供給學員）'),
]:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    br = p.add_run(f'• {name}：')
    set_rfont(br, 13, bold=True)
    nr = p.add_run(desc)
    set_rfont(nr, 13)

para('判斷邏輯：零件比對（精確）AND 動作比對（模糊關鍵詞）同時成立才算正確。'
     '答對不扣分，答錯扣 10 分；求助學長扣 5 分並給出方向性提示。')

h('6.2　五種故障情境 SOP 步驟定義', 2)
para('表 5 列出五種故障情境的 SOP 步驟摘要，每種情境包含 5 個有序步驟，'
     '步驟順序代表正確的工程處置流程。')
doc.add_paragraph()

tbl_sop = doc.add_table(rows=1, cols=4)
tbl_sop.style = 'Table Grid'
for cell, t in zip(tbl_sop.rows[0].cells,
                   ['故障類型', 'SOP 步驟（依序）', '觸發 Mesh', '相關子系統']):
    cell.text = t
    for r in cell.paragraphs[0].runs:
        set_rfont(r, 10, bold=True)

sop_rows = [
    ('鏡片熱點\nlens_hotspot',
     '①確認鏡片溫度 →②降低曝光劑量 →③停機等待冷卻 →④確認冷卻水迴路 →⑤恢復曝光監控 CDU',
     'POB_Barrel / Lens_0~9',
     '投影鏡組熱控'),
    ('光罩污染\ncontamination',
     '①CDU Map 確認 →②卸載光罩 →③目視檢查 →④清潔/更換 →⑤裝回執行對準',
     'Reticle_Stage / Reticle',
     '光罩載台系統'),
    ('載台誤差\nstage_error',
     '①確認 Overlay 超規 →②暫停曝光 →③執行基準校正 →④確認震動/溫度 →⑤驗證 Overlay 回規',
     'Wafer_Chuck / Stage_Base',
     '晶圓台定位'),
    ('劑量漂移\ndose_drift',
     '①確認 CD 偏移 →②查看劑量監控值 →③校正劑量設定 →④曝光驗證片 →⑤確認 CDU 回規',
     'Laser_Box / 控制系統',
     '雷射能量控制'),
    ('焦距漂移\nfocus_drift',
     '①確認焦距漂移量 →②執行 Focus 校正 →③查看鏡片溫升 →④等待熱穩定 →⑤恢復量產監控',
     'Illum_Barrel / 控制系統',
     '焦距熱補償'),
]
for ft, steps, mesh, sys in sop_rows:
    row = tbl_sop.add_row().cells
    row[0].text=ft; row[1].text=steps; row[2].text=mesh; row[3].text=sys
    for c in row:
        for r in c.paragraphs[0].runs:
            set_rfont(r, 10)
caption('表 5　五種故障情境 SOP 步驟定義摘要')
doc.add_paragraph()

h('6.3　自適應模式切換邏輯', 2)
para('評估系統依連續操作結果動態調整教學模式，共四個等級：')
for mode, cond, style in [
    ('挑戰模式（Challenge）', '連續答對 ≥ 3 次 → 從標準升級',   '回饋簡潔、無額外提示'),
    ('標準模式（Standard）',  '預設模式',                          '適量提示、平衡回饋'),
    ('鷹架模式（Scaffolding）','連續答錯 ≥ 3 次 → 從標準降級',  '頻繁提示、詳細說明'),
    ('補救模式（Remedial）',  '連續答錯 ≥ 5 次 → 從鷹架降級',  '提供 reason 原因解釋、超詳細引導'),
]:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    br = p.add_run(f'• {mode}：{cond}。')
    set_rfont(br, 13, bold=True)
    nr = p.add_run(f'  回饋風格：{style}')
    set_rfont(nr, 13)

para('此機制與 SECOM 資料集互相獨立：SECOM 負責感測器統計特性與故障健康分數，'
     'SOP 規則表負責操作序列的正確性驗證，兩者共同構成本系統的知識基礎。')

# ── 七、訓練系統流程 ──────────────────────────────────────────────────────────
h('七、訓練系統流程', 1)
para('本系統採整合式訓練設計，理論問答與 3D 實機操作於同一訓練情境中同步進行，'
     '如圖 13 所示。受訓者進入 3D 環境後即可自由操作設備、查看 HMI 面板，'
     '並隨時透過 AI 對話框提問；AI 導師根據輸入內容動態切換理論引導或實作指引，'
     '無需先通過任何獨立的理論測驗。')
doc.add_paragraph()
fig(fig_training(), w=6.2, cap='圖 13　訓練系統流程圖')

h('7.1　整合式訓練設計', 2)
para('受訓者進入整合訓練情境後，理論學習與實機操作同時展開。'
     '3D 環境中，受訓者以 WASD 鍵盤移動、[E] 鍵與設備 Mesh 互動、[C] 鍵開啟 AI 對話。'
     '當受訓者詢問 DUV 微影原理、製程參數或故障原因等理論問題時，'
     'AI 導師（Qwen LLM）即切換至蘇格拉底式引導模式，以提問引導其主動推理；'
     '當受訓者進行故障排除操作時，'
     '系統採「自主評估」而非「逐步指引」：受訓者自行決定靠近哪個零件、選擇何種操作，'
     'AI 即時評估並給予回饋，兩種模式在同一會話中無縫動態切換。')

h('7.2　故障情境與自主操作評估', 2)
para('系統共提供五種故障情境：鏡片熱點、光罩污染、'
     '晶圓台位置誤差、劑量漂移與焦距漂移。每種情境配有明確的觸發 Mesh 群組（故障發生時橘色發光提示），'
     '受訓者在 3D 場景中接近對應設備並按 [E] 鍵即可開啟零件互動面板。'
     '與傳統 SOP 逐步引導不同，本系統採「自主評估」設計：'
     '受訓者自行判斷正確操作順序，每次選擇操作後由 /api/action 端點'
     '對照 sop_definitions.py 中預定義的正確步驟序列進行比對，'
     '並依操作結果給予即時文字回饋（答對原因說明 / 答錯糾正提示）。'
     '連續答錯 3 次後系統自動降低難度，連續答對 3 次後升高難度（自適應模式切換）；'
     '若受訓者需要提示，可按「求助學長」按鈕（扣 5 分）獲得當前步驟的方向性提示。'
     '完成全部步驟後，訓練總評分達 80 分以上即取得能力認證。')

h('7.3　評分系統（ScoringSystem）', 2)
para('評分系統（ScoringSystem）依三個主要維度對每次訓練會話進行量化評估：'
     '（1）診斷準確性（Diagnostic Accuracy）佔 30%，評估故障類型判斷正確性及信心度加權；'
     '（2）操作正確性（Operation Correctness）佔 40%，評估 SOP 步驟符合度及操作順序；'
     '（3）安全合規性（Safety Compliance）佔 30%，評估 PPE 穿戴與安全規範遵守情況。'
     '時間效率（Time Efficiency）另設加分機制，以實際完成時間相對於標準時限之比率計算。'
     '最終總分依 A（≥90）、B（≥80）、C（≥70）、D（≥60）、F（<60）五個等級評定，'
     '達 B 級以上（≥80 分）即取得能力認證，並由系統自動推薦進階訓練模組。')

# ── 八、預期成果 ──────────────────────────────────────────────────────────────
doc.add_page_break()
h('八、預期成果與研究貢獻', 1)
para('本研究預期透過整合 AI 導師、數位孿生與沉浸式 3D 訓練環境，'
     '在以下五個面向達成具體成效。')
doc.add_paragraph()

h('8.1　降低人力訓練成本與機台停機損失', 2)
para('半導體微影機台屬高資本密度設備，ASML DUV 機台採購成本逾數千萬美元，'
     '設備停機一小時的機會成本估計可達數萬至數十萬美元（視晶圓廠產能與製程節點而定）。'
     '傳統新進工程師訓練流程通常需要至少 2–4 週實機操作練習，'
     '期間須由一名資深工程師全程陪同，佔用機台機時且消耗資深人力。')
para('更關鍵的風險在於：初學者在操作不熟悉的機台時，可能因步驟錯誤（如誤設曝光劑量、'
     '錯誤執行光罩裝卸程序、未正確處理真空系統異常）導致光罩損傷、晶圓刮傷，'
     '或引發設備聯鎖（interlock）觸發，造成批量晶圓報廢。'
     '以 300 mm 晶圓為例，一片晶圓的材料成本依製程節點不同約為數百至數千美元，'
     '若因訓練失誤導致整批報廢，損失往往難以估計。')
para('本虛擬訓練系統以瀏覽器端 WebGL 環境完全取代初期實機接觸：'
     '受訓者在零硬體風險的沉浸式 3D 環境中反覆練習五種故障排除情境，'
     '不佔用任何實機機時、無需資深工程師現場監護，'
     '預期可顯著壓縮達到上機資格所需的現場訓練時數，'
     '並降低因訓練失誤引發的 wafer loss 與機台故障風險。'
     '此外，系統可在任意時間、任意地點透過瀏覽器開啟，'
     '突破傳統「機台有空才能訓練」的時間與地點限制，提升訓練資源可及性。')

h('8.2　數位孿生驅動之高擬真訓練情境', 2)
para('本研究的核心技術特色在於以「設備數位孿生」（Lithography Digital Twin）作為訓練情境的物理基礎。'
     '數位孿生並非靜態的說明圖鑑，而是一個能即時反映設備狀態、模擬感測器讀值並響應操作行為的動態模型。')
para('具體而言，本系統整合 UCI SECOM 真實製程資料集（1,567 筆晶圓量測記錄、590 個感測器特徵）'
     '建立統計雜訊模型，使虛擬 HMI 面板顯示的溫度、流量、對準誤差、dose 讀值等感測器數值，'
     '均符合真實 DUV 設備在正常及異常狀態下的統計分布特性，而非固定數值或隨機假數據。'
     '配合基於 Bossung Curve 的焦距–劑量–CD 物理方程式與雙指數鏡片熱模型（lens heating model），'
     '系統能動態模擬鏡片過熱、dose 漂移、焦距偏移對 CDU Map 與 Overlay 的實際影響，'
     '讓受訓者在虛擬環境中觀察到與實機一致的製程退化現象。')
para('數位孿生的另一項重要優勢在於可安全地模擬「稀有但高危」的故障情境。'
     '在實際生產線上，光罩污染（contamination）、晶圓台位置誤差（stage error）等嚴重故障'
     '發生頻率低，一旦發生往往已造成大量晶圓報廢；新進工程師幾乎不可能在訓練期間遇到足夠次數。'
     '本系統透過數位孿生將這些情境轉化為可隨時重現、難度可調的標準訓練課題，'
     '使受訓者在無實際損失風險的條件下累積故障處理經驗，')
para('預期此數位孿生架構的導入，將成為半導體設備教育從「知識灌輸」走向「情境模擬與能力驗證」'
     '的關鍵技術橋梁，並為後續多設備、多製程的虛擬訓練平台建立可複製的技術基礎。')

h('8.3　個人化自適應學習成效', 2)
para('系統依據 AI 評分（0–10 分）即時切換挑戰、標準、鷹架、補救四種教學模式，'
     '避免「一體適用」課程導致程度差異被忽視的問題。'
     '預期對不同背景的受訓者（初學者至有經驗技術員）均能提供適配的引導強度，'
     '縮短達到操作合格門檻（≥80 分）所需之平均訓練次數。'
     '系統採用蘇格拉底式對話引導受訓者主動建構知識，而非被動接收指令；'
     '「求助學長」機制（-5 分扣點）在提供安全網的同時，保留主動思考的激勵誘因。'
     '對於已有實務經驗的資深技術員，可直接挑戰較高難度情境並跳過基礎鷹架引導，'
     '避免資深人員因訓練內容過於基礎而降低參與意願。')

h('8.4　客觀量化能力認證', 2)
para('傳統訓練成效評估多依賴主管主觀判斷或紙筆測驗，難以全面衡量實際操作能力。'
     '本系統以三維度加權評分（診斷準確性 30%、操作正確性 40%、安全合規 30%）'
     '搭配時間效率加分項，自動輸出量化分數與 A–F 五級制等級認定，'
     '提供可追蹤、可比較的客觀能力認證依據。'
     '學習歷程全程記錄（每次操作步驟、答對/答錯次數、求助次數、完成時間），'
     '可供管理者分析整體訓練品質趨勢，識別特定故障類型之高失誤率步驟，'
     '進而優化 SOP 設計或補強對應訓練課題。')

h('8.5　學術貢獻與可移植性', 2)
para('本研究作為數位孿生應用於半導體操作員訓練的先導實作，預期提供以下學術貢獻：'
     '（1）建立以公開資料集（UCI SECOM）驅動設備數位孿生感測器的方法論，'
     '示範從原始製程資料到可互動訓練情境的完整技術路徑；'
     '（2）驗證本地端 LLM（Qwen + Ollama）結合製程領域知識庫之 AI 導師可行性，'
     '無需連線外部 API 即可提供高品質蘇格拉底式引導，適合資訊安全要求嚴格的晶圓廠環境；'
     '（3）提出 Three.js WebGL 為核心的設備 3D 互動建模工作流程，'
     '供半導體教育領域數位化訓練系統參考與移植；'
     '（4）為後續研究奠定基礎，包括：多設備擴充（ALD、PVD、CMP 製程設備）、'
     '多人協同訓練模式，以及將數位孿生延伸至計量量測（Metrology）與良率管理（YMS）的完整工作流程。')
doc.add_paragraph()

# ── 九、結論 ──────────────────────────────────────────────────────────────────
doc.add_page_break()
h('九、結論', 1)
para('本研究提出一套以「設備數位孿生」為基礎、以「AI 導師」為引導核心之半導體微影設備虛擬訓練系統。'
     '系統透過沉浸式第一人稱 3D 環境，讓受訓者在不接觸實體設備、不佔用機台機時的情況下，'
     '於整合式訓練情境中同步進行理論問答與實機操作模擬，AI 導師動態切換引導模式。')
para('數位孿生的核心價值在於：以 UCI SECOM 真實製程資料集（590 個感測器特徵）建立統計雜訊模型，'
     '結合鏡片熱方程式與 Bossung Curve 物理模型，使虛擬 HMI 面板顯示的感測器讀值符合實機統計特性，'
     '讓受訓者在虛擬環境中觀察到與真實設備一致的製程退化現象與故障徵兆。'
     '這一特性使系統得以安全重現光罩污染、晶圓台位置誤差等高危故障情境，'
     '讓新進操作員在零 wafer 損失風險下累積足夠的故障處理經驗，'
     '預期可顯著降低因訓練不足導致的實機操作失誤、晶圓報廢與機台停機損失。')
para('多代理 AI 架構（Diagnostic Agent、Operation Agent、Adaptive Teaching）實現個人化四模式導師引導；'
     '三維度量化評分系統提供客觀、可量化的能力認證依據。'
     '後續研究方向包括：引入實際訓練成效量化比較（受訓前後診斷正確率、操作完成時間）、'
     '擴充至多機台設備訓練模組（ALD、CMP、Metrology），'
     '以及探討多人協同訓練模式，進一步驗證數位孿生在半導體人才培育中的實際效益。')

# ── 十、名詞說明 ──────────────────────────────────────────────────────────────
h('十、名詞說明（Glossary）', 1)
terms = [
    ('DUV（Deep Ultraviolet，深紫外光）',
     '波長範圍 100–300 nm 之紫外光。本系統採用 248 nm ArF 準分子雷射，'
     '用於曝光光阻以定義電路圖案。波長越短，理論解析度越高（Rayleigh 準則：R = kλ/NA）。'),
    ('ArF 準分子雷射（ArF Excimer Laser）',
     '氬氟（ArF）準分子雷射，輸出波長 193 nm（本系統以 248 nm KrF 為代表）。'
     '提供 DUV 微影所需之高功率脈衝雷射光源。'),
    ('CD（Critical Dimension，臨界尺寸）',
     '光罩圖案轉移至晶圓後，最小特徵結構的線寬，單位為奈米（nm）。'
     'CD 控制是微影製程品質的核心指標，直接影響電晶體性能與電路良率。'),
    ('CDU（CD Uniformity，線寬均勻性）',
     '整片晶圓或單一曝光場內 CD 分布的均勻程度，通常以 3σ 表示。'
     'CDU 受焦距均勻性、照明均勻性、光阻批次差異及鏡頭像差等因素影響。'),
    ('Overlay（疊對誤差）',
     '不同微影層之間圖案對準的偏差量，單位為奈米。'
     'Overlay 誤差來源包括晶圓台定位誤差、鏡頭彗差（Zernike Z7/Z8）及熱膨脹。'
     '先進製程節點要求 Overlay < 2 nm。'),
    ('NA（Numerical Aperture，數值孔徑）',
     '描述光學系統收光能力的無因次參數，NA = n·sinθ。NA 越大，解析度越高，'
     '但焦深（DOF）越小。現代 DUV 浸潤式機台 NA 可達 1.35（以水為介質）。'),
    ('DOF（Depth of Focus，焦深）',
     '在維持可接受 CD 品質的前提下，焦平面可偏移的範圍。'
     'DOF ∝ λ/NA²，NA 增加會快速縮減焦深，對製程控制造成挑戰。'),
    ('SECOM（Semiconductor Equipment COmputer Monitoring）',
     'UCI 機器學習資料庫中的半導體製程監控資料集（2008 年）。'
     '包含 1,567 筆晶圓量測記錄與 590 個製程感測器特徵，標記為通過/失敗。'
     '本系統用以建立真實統計基礎的製程雜訊模型。'),
    ('PCA（Principal Component Analysis，主成分分析）',
     '一種線性降維方法，將高維特徵空間投影至方差最大的正交主成分軸上。'
     '本系統對 SECOM 590 個特徵進行 PCA，萃取前 20 個最具判別力的主成分，'
     '用以計算設備健康分數。'),
    ('HMI（Human-Machine Interface，人機介面）',
     '操作人員與設備之間的互動介面，通常為觸控式顯示螢幕。'
     '本系統以 HTML5 Canvas 渲染虛擬 HMI，呈現即時感測器讀值、警報訊息與製程控制面板。'),
    ('APC（Advanced Process Control，先進製程控制）',
     '以即時量測回饋（Run-to-Run control）自動調整製程參數，'
     '維持 CD/Overlay 在規格範圍內，是現代半導體製程品質管控的核心機制。'),
    ('SOP（Standard Operating Procedure，標準作業程序）',
     '針對特定故障類型的標準化步驟化處置流程，確保維修操作的一致性與安全性。'
     '本系統提供 5 種故障類型的 SOP，每種包含 5 個執行步驟。'),
    ('PointerLock API',
     'Web 瀏覽器 API，允許應用程式鎖定滑鼠游標並持續讀取相對移動量，'
     '實現第一人稱視角控制而不受螢幕邊界限制，常見於 3D 遊戲應用。'),
    ('GLB（GL Transmission Format Binary）',
     'glTF 格式的二進位版本，將 3D 幾何、材質與場景結構封裝於單一檔案中。'
     '本系統之 ASML DUV 設備模型以 asml_duv.glb 格式提供，透過 Three.js GLTFLoader 載入。'),
    ('Raycaster（射線投射）',
     '三維場景中，從相機位置沿視線方向投射射線，偵測與場景物件（Mesh）交點的演算法。'
     '本系統用以實現「準心對準物件 → 按 E 互動」的第一人稱交互功能。'),
    ('Qwen（通義千問）',
     '阿里巴巴開發之大型語言模型（LLM），支援本地端推理。'
     '本系統透過 Ollama 或 llama.cpp 後端，在無需連網的情況下提供 AI 導師對話功能。'),
    ('A2A（Agent-to-Agent）',
     '多代理系統中，代理之間直接溝通與協調的架構模式。'
     '本系統以 A2ACoordinator 統一管理診斷、操作、安全及教學等多個 AI 代理。'),
]
for term, defn in terms:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.3)
    p.paragraph_format.space_before = Pt(4)
    br = p.add_run(term)
    set_rfont(br, 13, bold=True)
    br.font.color.rgb = RGBColor(0x1A,0x3A,0x5C)
    nr = p.add_run('\n　　' + defn)
    set_rfont(nr, 13)

# ── 參考文獻 ──────────────────────────────────────────────────────────────────
h('參考文獻', 1)
refs = [
    '[1] ASML. (2023). TWINSCAN NXT:2050i Product Specifications. ASML Technical Documentation.',
    '[2] McCord, M.A. & Rooks, M.J. (2000). Handbook of Microlithography, Micromachining, and Microfabrication. SPIE Press.',
    '[3] Mack, C. (2007). Fundamental Principles of Optical Lithography. Wiley-Interscience.',
    '[4] Smith, B.W. (2007). Microlithography: Science and Technology, 2nd ed. CRC Press.',
    '[5] UCI Machine Learning Repository. (2008). SECOM Dataset. UC Irvine.',
    '[6] Zernike, F. (1934). Diffraction theory of the knife-edge test. Monthly Notices RAS, 94, 377–384.',
    '[7] Bossung, J.W. (1977). Projection printing characterization. SPIE, 100, 80–84.',
    '[8] Three.js Documentation. (2024). https://threejs.org/docs/',
    '[9] 張，等. (2023). 半導體製程技術概論. 全華圖書.',
]
for ref in refs:
    rp = doc.add_paragraph(ref)
    rp.paragraph_format.left_indent = Cm(0.5)
    rp.paragraph_format.first_line_indent = Cm(-0.5)
    for r in rp.runs:
        set_rfont(r, 13)

# ── 儲存 ──────────────────────────────────────────────────────────────────────
OUT = r"c:\Users\user\Desktop\在職碩\OneDrive - 長庚大學\長庚碩班\論文\系統架構說明書.docx"
doc.save(OUT)
print(f"Saved: {OUT}")
